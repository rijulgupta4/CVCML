from __future__ import annotations

import json
import shutil
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r"C:\path\to\CVCML")
OUTPUTS = PROJECT / "Outputs"
OUT = OUTPUTS / "Run 28 (v0.5 Model Card and Manuscript Package)"
FIGURES = OUT / "figures"

RUN16 = OUTPUTS / "Run 16 (v0.5 Catheter Episode Redesign)"
RUN17 = OUTPUTS / "Run 17 (v0.5 Daily Landmark Frame)"
RUN23 = OUTPUTS / "Run 23 (v0.5 Label Sensitivity Modeling)"
RUN24 = OUTPUTS / "Run 24 (v0.5 Operating Policy Characterization)"
RUN25 = OUTPUTS / "Run 25 (v0.5 Locked Temporal Evaluation)"
RUN26 = OUTPUTS / "Run 26 (v0.5 Locked Error Analysis)"
RUN27 = OUTPUTS / "Run 27 (v0.5 ICD Agreement Label Validation)"
RUN271 = OUTPUTS / "Run 27.1 (v0.5 ICD Discordance Supplement)"

BLUE = "2E74B5"
NAVY = "17365D"
DARK = "1F2937"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_RED = "FCE8E6"
PALE_GREEN = "E8F3EC"
WHITE = "FFFFFF"

CDC_URL = "https://www.cdc.gov/nhsn/pdfs/pscmanual/4psc_clabscurrent.pdf"
MIMIC_DIAG_URL = "https://mimic.mit.edu/docs/IV/modules/hosp/diagnoses_icd.html"
MIMIC_PROC_URL = "https://mimic.mit.edu/docs/IV/modules/icu/procedureevents.html"
TRIPOD_URL = "https://www.bmj.com/content/385/bmj-2023-078378"
PROBAST_URL = "https://www.bmj.com/content/388/bmj-2024-082505"
ALBU_URL = "https://pubmed.ncbi.nlm.nih.gov/40275180/"
RAHMANI_URL = "https://pubmed.ncbi.nlm.nih.gov/34428529/"


def read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(path)
    return pd.read_csv(path)


def percent(value: float, digits: int = 1) -> str:
    return f"{100 * float(value):.{digits}f}%"


def metric(value: float, digits: int = 3) -> str:
    return f"{float(value):.{digits}f}"


def clean_feature(name: str) -> str:
    return name.replace("numeric__", "").replace("categorical__", "")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_document(doc: Document, preset: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    if preset == "compact_reference_guide":
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        heading_specs = {
            "Title": (26, NAVY, 0, 16),
            "Heading 1": (16, BLUE, 18, 10),
            "Heading 2": (13, BLUE, 14, 7),
            "Heading 3": (12, NAVY, 10, 5),
        }
    else:
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.333
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        heading_specs = {
            "Title": (28, NAVY, 0, 18),
            "Heading 1": (17, BLUE, 18, 10),
            "Heading 2": (14, BLUE, 12, 6),
            "Heading 3": (12, NAVY, 8, 4),
        }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    header = section.header.paragraphs[0]
    header.text = "CVCML v0.5  |  Run 28 evidence package"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Research prototype - not for clinical use  |  Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(footer)


def add_title_block(doc: Document, title: str, subtitle: str, status: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(title)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=8, line=1.1)
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=18)
    run = p.add_run(status)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(BLUE)


def add_callout(doc: Document, heading: str, body: str, fill=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, start=180, bottom=120, end=180)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=4)
    run = p.add_run(heading)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p = cell.add_paragraph(body)
    set_paragraph_spacing(p, after=0, line=1.15)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], header_fill=LIGHT_BLUE, font_size=9) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(NAVY)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            cell.text = str(value)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=0, line=1.0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)


def add_figure(doc: Document, path: Path, caption: str, width=6.15) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_reference_paragraph(doc: Document, number: int, text: str, url: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=5, line=1.05)
    lead = p.add_run(f"{number}. ")
    lead.bold = True
    p.add_run(text + " ")
    link = p.add_run(url)
    link.font.color.rgb = RGBColor.from_string(BLUE)
    link.font.underline = True


def load_evidence() -> dict:
    episode_audit = read_csv(RUN16 / "v0_5_episode_label_audit.csv").iloc[0]
    temporal = read_csv(RUN17 / "v0_5_daily_landmark_temporal_split_audit.csv")
    model23 = read_csv(RUN23 / "v0_5_run23_label_sensitivity_model_comparison.csv")
    split23 = read_csv(RUN23 / "v0_5_run23_label_sensitivity_split_audit.csv")
    importance = read_csv(RUN23 / "v0_5_run23_label_sensitivity_feature_importance.csv")
    lockbox = read_csv(RUN25 / "v0_5_run25_lockbox_model_comparison.csv")
    val_lock = read_csv(RUN25 / "v0_5_run25_validation_lockbox_comparison.csv")
    row_policy = read_csv(RUN25 / "v0_5_run25_lockbox_topk_row_policy.csv")
    episode_policy = read_csv(RUN25 / "v0_5_run25_lockbox_topk_episode_policy.csv")
    feature_contrast = read_csv(RUN26 / "v0_5_run26_feature_error_contrast.csv")
    icd = read_csv(RUN27 / "v0_5_run27_proxy_icd_agreement_table.csv")
    grain = read_csv(RUN271 / "v0_5_run27_1_grain_reconciled_agreement.csv")

    primary_val = model23[(model23.target_label == "primary_or_uncertain") & (model23.split == "validation") & (model23.calibration == "platt")].iloc[0]
    primary_lock = lockbox[lockbox.calibration == "platt"].iloc[0]
    row5 = row_policy[row_policy.policy == "top_5%_rows"].iloc[0]
    row10 = row_policy[row_policy.policy == "top_10%_rows"].iloc[0]
    ep100 = episode_policy[episode_policy.policy == "top_100_episodes"].iloc[0]
    ep150 = episode_policy[episode_policy.policy == "top_150_episodes"].iloc[0]
    ep250 = episode_policy[episode_policy.policy == "top_250_episodes"].iloc[0]
    icd_primary = icd[(icd.group == "all") & (icd.proxy_label == "cvc_bsi_strict_primary_or_uncertain_proxy") & (icd.icd_comparator == "icd_cvc_bsi_specific")].iloc[0]
    grain_episode = grain[grain.grain == "catheter_episode"].iloc[0]
    grain_hadm = grain[grain.grain == "hospital_admission"].iloc[0]
    top_importance = importance[importance.target_label == "primary_or_uncertain"].nlargest(15, "importance").copy()

    return {
        "episode_audit": episode_audit,
        "temporal": temporal,
        "model23": model23,
        "split23": split23,
        "primary_val": primary_val,
        "primary_lock": primary_lock,
        "val_lock": val_lock,
        "row5": row5,
        "row10": row10,
        "ep100": ep100,
        "ep150": ep150,
        "ep250": ep250,
        "feature_contrast": feature_contrast,
        "icd_primary": icd_primary,
        "grain_episode": grain_episode,
        "grain_hadm": grain_hadm,
        "top_importance": top_importance,
    }


def copy_figures() -> dict[str, Path]:
    FIGURES.mkdir(parents=True, exist_ok=True)
    mapping = {
        "validation_lockbox": RUN25 / "plots" / "v0_5_run25_validation_vs_lockbox.png",
        "calibration": RUN25 / "plots" / "v0_5_run25_lockbox_calibration_deciles.png",
        "error_contrast": RUN26 / "plots" / "v0_5_run26_feature_contrast_fp_vs_fn.png",
        "icd_overlap": RUN27 / "plots" / "v0_5_run27_proxy_icd_overlap.png",
        "icd_era": RUN271 / "plots" / "v0_5_run27_1_discordance_by_era.png",
    }
    result = {}
    for key, source in mapping.items():
        if not source.exists():
            raise FileNotFoundError(source)
        target = FIGURES / f"v0_5_run28_{key}.png"
        shutil.copy2(source, target)
        result[key] = target
    return result


def write_consolidated_tables(e: dict) -> dict[str, Path]:
    OUT.mkdir(parents=True, exist_ok=True)
    a = e["episode_audit"]
    cohort_rows = [
        ["Raw CVC procedure events", int(a.raw_cvc_procedure_events), "Run 16 episode audit"],
        ["Continuous exposure periods", int(a.continuous_exposure_periods), "Run 16 episode audit"],
        ["Eligible exposure periods (>=48 h)", int(a.eligible_48h_exposure_periods), "Run 16 episode audit"],
        ["Stays with multiple exposure periods", int(a.stays_with_multiple_exposure_periods), "Run 16 episode audit"],
        ["Broad proxy-positive eligible episodes", int(a.broad_proxy_positive_episodes), "Run 16 episode audit"],
        ["Strict proxy-positive eligible episodes", int(a.strict_proxy_positive_episodes), "Run 16 episode audit"],
        ["Daily landmark rows before later source screening", int(e["temporal"].rows.sum()), "Run 17 temporal audit"],
        ["Scored lockbox rows", int(e["primary_lock"].rows), "Run 25 locked evaluation"],
        ["Scored lockbox episodes", int(e["primary_lock"].episodes), "Run 25 locked evaluation"],
        ["Scored lockbox patients", int(e["primary_lock"].patients), "Run 25 locked evaluation"],
    ]
    cohort = pd.DataFrame(cohort_rows, columns=["measure", "value", "source"])

    performance = pd.DataFrame([
        ["Validation", "2017-2019", e["primary_val"].rows, None, None, e["primary_val"].positive_rows, None, e["primary_val"].prevalence, e["primary_val"].roc_auc, e["primary_val"].pr_auc, e["primary_val"].pr_auc_lift_over_prevalence, e["primary_val"].brier_skill_score, e["primary_val"].expected_observed_ratio, None, None],
        ["Temporal lockbox", "2020-2022", e["primary_lock"].rows, e["primary_lock"].episodes, e["primary_lock"].patients, e["primary_lock"].positive_rows, e["primary_lock"].positive_episodes, e["primary_lock"].prevalence, e["primary_lock"].roc_auc, e["primary_lock"].pr_auc, e["primary_lock"].pr_auc_lift_over_prevalence, e["primary_lock"].brier_skill_score, e["primary_lock"].expected_observed_ratio, e["primary_lock"].calibration_intercept, e["primary_lock"].calibration_slope],
    ], columns=["split", "years", "rows", "episodes", "patients", "positive_rows", "positive_episodes", "prevalence", "roc_auc", "pr_auc", "pr_auc_lift", "brier_skill_score", "expected_observed_ratio", "calibration_intercept", "calibration_slope"])

    def policy_row(name, p):
        return [name, p.policy_family, int(p.rows_reviewed), int(p.episodes_reviewed), p.precision_ppv, p.ppv_lift_over_prevalence, p.row_recall_sensitivity, int(p.positive_episodes_captured), p.episode_recall_sensitivity, p.false_reviews_per_true_positive]
    policies = pd.DataFrame([
        policy_row("Top 5% landmark rows", e["row5"]),
        policy_row("Top 10% landmark rows", e["row10"]),
        policy_row("Top 100 episodes", e["ep100"]),
        policy_row("Top 150 episodes", e["ep150"]),
        policy_row("Top 250 episodes", e["ep250"]),
    ], columns=["policy", "policy_family", "reviews", "episodes_reviewed", "ppv", "ppv_lift", "row_recall", "positive_episodes_captured", "episode_recall", "false_reviews_per_true_positive"])

    limitations = pd.DataFrame([
        ["Outcome validity", "Primary outcome is a strict CVC-associated BSI proxy, not adjudicated NHSN CLABSI.", "Retain proxy terminology; adjudicate a sample using notes and source rules.", "CDC NHSN manual; Runs 22, 27"],
        ["Episode ascertainment", "procedureevents presence is positive evidence, but missing records do not prove absence of a line.", "Validate line exposure against another database or chart-derived source.", "MIMIC procedureevents documentation"],
        ["Outcome-adjacent feature", "early_positive_culture is included among candidate predictors and may encode information too close to the endpoint.", "Run a focused ablation and timestamp audit before any clinical claim.", "Run 23 feature inventory"],
        ["Temporal generalization", "The lockbox is later in time but remains from one institution.", "Perform external validation on a second database.", "Run 25"],
        ["Precision", "The lockbox contains only 26 positive episodes and no final cluster-bootstrap confidence intervals.", "Add episode/patient-cluster bootstrap intervals without model redevelopment.", "Run 25"],
        ["Operational burden", "Even the best review-list policies produce several false reviews for each true-positive row.", "Frame as infection-prevention prioritization, not an interruptive bedside alert.", "Runs 24-26"],
        ["Fairness", "Subgroup stability has not been fully established for the frozen v0.5 model.", "Report performance and calibration by prespecified demographic and care-setting groups.", "Open item"],
        ["ICD comparison", "ICD coding is admission-level and discharge-derived; episode-level overlap is expected to be low.", "Use ICD agreement as convergent validity only, not ground truth.", "MIMIC diagnoses documentation; Runs 27-27.1"],
    ], columns=["domain", "current_limitation", "required_action", "evidence"])

    evidence = pd.DataFrame([
        ["Catheter episode reconstruction", "Run 16", str(RUN16 / "v0_5_episode_label_audit.csv")],
        ["Daily landmark and temporal frame", "Run 17", str(RUN17 / "v0_5_daily_landmark_temporal_split_audit.csv")],
        ["Label-sensitive validation modeling", "Run 23", str(RUN23 / "v0_5_run23_label_sensitivity_model_comparison.csv")],
        ["Operating policy development", "Run 24", str(RUN24 / "v0_5_run24_operating_policy_summary.csv")],
        ["Frozen temporal lockbox evaluation", "Run 25", str(RUN25 / "v0_5_run25_lockbox_model_comparison.csv")],
        ["Locked error analysis", "Run 26", str(RUN26 / "v0_5_run26_feature_error_contrast.csv")],
        ["ICD-coded agreement", "Run 27", str(RUN27 / "v0_5_run27_proxy_icd_agreement_table.csv")],
        ["Grain-reconciled discordance supplement", "Run 27.1", str(RUN271 / "v0_5_run27_1_grain_reconciled_agreement.csv")],
    ], columns=["claim_area", "run", "source_file"])

    paths = {
        "cohort": OUT / "v0_5_run28_cohort_summary.csv",
        "performance": OUT / "v0_5_run28_performance_summary.csv",
        "policies": OUT / "v0_5_run28_operating_policy_summary.csv",
        "limitations": OUT / "v0_5_run28_limitations_and_actions.csv",
        "evidence": OUT / "v0_5_run28_evidence_traceability.csv",
    }
    cohort.to_csv(paths["cohort"], index=False)
    performance.to_csv(paths["performance"], index=False)
    policies.to_csv(paths["policies"], index=False)
    limitations.to_csv(paths["limitations"], index=False)
    evidence.to_csv(paths["evidence"], index=False)
    return paths


def build_model_card(e: dict, figures: dict[str, Path], table_paths: dict[str, Path]) -> Path:
    doc = Document()
    configure_document(doc, "compact_reference_guide")
    add_title_block(
        doc,
        "CVCML v0.5 Model Card",
        "Seven-day strict CVC-associated BSI proxy risk prioritization from MIMIC-IV",
        f"Run 28 | Frozen evidence through Run 27.1 | {date.today().isoformat()}",
    )
    add_callout(
        doc,
        "Intended use",
        "Research-grade infection-prevention review-list prioritization. The model is not validated for autonomous diagnosis, bedside interruption, line-removal decisions, or direct patient care.",
        PALE_GREEN,
    )
    add_callout(
        doc,
        "Outcome language",
        "The target is a strict primary-or-uncertain CVC-associated bloodstream infection proxy within seven days. It is not an adjudicated NHSN CLABSI outcome.",
        PALE_RED,
    )

    doc.add_heading("Model At A Glance", level=1)
    add_table(doc, ["Field", "Frozen specification"], [
        ["Model", "XGBoost classifier followed by Platt calibration"],
        ["Prediction unit", "One daily landmark within an eligible catheter-exposure episode"],
        ["Prediction horizon", "Strict primary-or-uncertain CVC-associated BSI proxy in the next 7 days"],
        ["Data source", "MIMIC-IV v3.1 ICU and hospital modules"],
        ["Development period", "2008-2013; calibration 2014-2016; validation 2017-2019"],
        ["Locked temporal test", "2020-2022; opened once in Run 25"],
        ["Feature families", "Static context, labs, vitals, therapy context, line-exposure timing"],
        ["Candidate inputs", "173 pre-encoding fields (168 numeric, 5 categorical)"],
        ["Status", "Research prototype; no clinical deployment authorization"],
    ], [2700, 6660])

    doc.add_heading("Cohort And Label", level=1)
    a = e["episode_audit"]
    add_table(doc, ["Measure", "Value", "Interpretation"], [
        ["Raw CVC procedure records", f"{int(a.raw_cvc_procedure_events):,}", "Positive evidence of recorded catheter procedures"],
        ["Continuous exposure periods", f"{int(a.continuous_exposure_periods):,}", "All reconstructed periods, retaining multiple periods per stay"],
        ["Eligible periods", f"{int(a.eligible_48h_exposure_periods):,}", ">=48 hours of recorded exposure"],
        ["Broad proxy positives", f"{int(a.broad_proxy_positive_episodes):,} ({percent(a.broad_proxy_positive_rate)})", "Sensitivity outcome"],
        ["Strict organism-rule positives", f"{int(a.strict_proxy_positive_episodes):,} ({percent(a.strict_proxy_positive_rate)})", "Pre-source-screen strict proxy"],
        ["Primary source-screened positives", "291 episodes (1.28% of all reconstructed episodes)", "Primary-or-uncertain proxy used in the final model"],
    ], [3400, 1900, 4060])
    p = doc.add_paragraph("The label requires reconstructed catheter exposure and qualifying blood-culture timing and organism logic, then separates primary-likely, secondary-possible, and uncertain source classes. It does not fully reconstruct NHSN secondary-source, symptom, or MBI-LCBI adjudication rules.")
    p.paragraph_format.keep_with_next = False

    doc.add_heading("Frozen Temporal Performance", level=1)
    val = e["primary_val"]
    lock = e["primary_lock"]
    add_table(doc, ["Metric", "Validation 2017-2019", "Lockbox 2020-2022"], [
        ["Rows / positive rows", f"{int(val.rows):,} / {int(val.positive_rows):,}", f"{int(lock.rows):,} / {int(lock.positive_rows):,}"],
        ["Positive episode count", "54", f"{int(lock.positive_episodes):,}"],
        ["Prevalence", percent(val.prevalence, 2), percent(lock.prevalence, 2)],
        ["ROC-AUC", metric(val.roc_auc), metric(lock.roc_auc)],
        ["PR-AUC", metric(val.pr_auc), metric(lock.pr_auc)],
        ["PR-AUC lift", f"{metric(val.pr_auc_lift_over_prevalence, 2)}x", f"{metric(lock.pr_auc_lift_over_prevalence, 2)}x"],
        ["Brier Skill Score", metric(val.brier_skill_score), metric(lock.brier_skill_score)],
        ["Expected:observed", metric(val.expected_observed_ratio), metric(lock.expected_observed_ratio)],
        ["Calibration intercept / slope", "Not reported: numerically unstable validation estimate", f"{metric(lock.calibration_intercept)} / {metric(lock.calibration_slope)}"],
    ], [3300, 3030, 3030])
    add_figure(doc, figures["validation_lockbox"], "Figure 1. Frozen validation-to-lockbox comparison. Discrimination remained modest, while PR-AUC increased with the higher lockbox prevalence.")
    add_figure(doc, figures["calibration"], "Figure 2. Platt-calibrated lockbox reliability by risk decile. The overall slope was 1.056 and E:O was 0.888, but the small positive count limits precision.")

    doc.add_heading("Operating Characteristics", level=1)
    add_callout(doc, "Recommended framing", "Use the score to rank a bounded daily review list for infection-prevention staff. Do not frame it as an interruptive nurse alarm.")
    row5, row10, ep150, ep250 = e["row5"], e["row10"], e["ep150"], e["ep250"]
    add_table(doc, ["Locked policy", "Reviews", "PPV", "Episode recall", "False reviews / TP"], [
        ["Top 5% landmark rows", f"{int(row5.rows_reviewed)} rows", percent(row5.precision_ppv), percent(row5.episode_recall_sensitivity), metric(row5.false_reviews_per_true_positive, 1)],
        ["Top 10% landmark rows", f"{int(row10.rows_reviewed)} rows", percent(row10.precision_ppv), percent(row10.episode_recall_sensitivity), metric(row10.false_reviews_per_true_positive, 1)],
        ["Top 150 episodes", "150 episodes", percent(ep150.precision_ppv), percent(ep150.episode_recall_sensitivity), metric(ep150.false_reviews_per_true_positive, 1)],
        ["Top 250 episodes", "250 episodes", percent(ep250.precision_ppv), percent(ep250.episode_recall_sensitivity), metric(ep250.false_reviews_per_true_positive, 1)],
    ], [2900, 1600, 1400, 1700, 1760])
    p = doc.add_paragraph("Primary policy values above come directly from Run 25's target-aligned lockbox evaluation. Run 26's post-hoc episode capture analysis used a broader episode-level estimand and is exploratory, so its higher capture rates must not replace these prespecified values.")
    p.runs[0].italic = True

    doc.add_heading("Model Behavior", level=1)
    top = e["top_importance"]
    add_table(doc, ["Rank", "Feature", "Gain importance"], [[str(j + 1), clean_feature(r.feature), metric(r.importance, 4)] for j, (_, r) in enumerate(top.head(10).iterrows())], [900, 6500, 1960])
    add_figure(doc, figures["error_contrast"], "Figure 3. Locked error-analysis contrasts. High-risk false positives often had greater measurement intensity and treatment context than missed positives.")
    add_callout(doc, "Important audit item", "early_positive_culture appears among the candidate predictors and is outcome-adjacent. A timestamp audit and exclusion ablation are required before any stronger causal, clinical, or deployment claim.", PALE_RED)

    doc.add_heading("Label Agreement", level=1)
    icd = e["icd_primary"]
    gh = e["grain_hadm"]
    add_table(doc, ["Comparison", "Proxy positive", "ICD positive", "Both", "Positive-set Jaccard"], [
        ["Catheter episode", f"{int(icd.proxy_positive):,}", f"{int(icd.icd_positive):,}", f"{int(icd.both_positive):,}", percent(icd.jaccard_positive, 1)],
        ["Hospital admission", f"{int(gh.proxy_positive):,}", f"{int(gh.icd_positive):,}", f"{int(gh.both_positive):,}", percent(gh.positive_set_jaccard, 1)],
    ], [2600, 1700, 1700, 1400, 1960])
    p = doc.add_paragraph("Low overlap does not establish proxy failure because the comparators differ in timing, grain, and purpose: the model label is episode- and time-specific, whereas MIMIC ICD diagnoses are admission-level codes assigned from the completed hospitalization record. ICD agreement is therefore convergent validity, not ground truth.")
    add_figure(doc, figures["icd_overlap"], "Figure 4. Episode-level overlap between the strict primary-or-uncertain proxy and specific ICD-coded catheter-related BSI.")

    doc.add_heading("Limitations And Required Safeguards", level=1)
    limits = read_csv(table_paths["limitations"])
    add_table(doc, ["Domain", "Current limitation", "Required action"], [[r.domain, r.current_limitation, r.required_action] for _, r in limits.iterrows()], [1900, 3760, 3700], font_size=8)

    doc.add_heading("Decision", level=1)
    add_callout(doc, "Current conclusion", "The frozen v0.5 model demonstrates modest temporal risk stratification and usable prioritization lift, with encouraging aggregate calibration after Platt scaling. It is credible as a research and portfolio artifact, but not yet as a deployable CLABSI alarm. The next decision should be based on whether the project goal is a manuscript draft or a focused validity improvement.", PALE_GREEN)
    add_table(doc, ["If the goal is...", "Next action"], [
        ["Portfolio/manuscript draft", "Proceed with the frozen package, add confidence intervals and subgroup reporting, and preserve the 2020-2022 lockbox."],
        ["Stronger clinical validity", "Perform one focused label improvement: notes/source adjudication plus the early_positive_culture exclusion audit."],
        ["Transportability", "Validate the frozen pipeline on a second database before any further MIMIC-based tuning."],
    ], [2800, 6560])

    doc.add_heading("References", level=1)
    add_reference_paragraph(doc, 1, "CDC/NHSN Patient Safety Component Manual, Bloodstream Infection Event chapter (January 2026).", CDC_URL)
    add_reference_paragraph(doc, 2, "MIMIC-IV diagnoses_icd module documentation.", MIMIC_DIAG_URL)
    add_reference_paragraph(doc, 3, "MIMIC-IV procedureevents module documentation.", MIMIC_PROC_URL)
    add_reference_paragraph(doc, 4, "TRIPOD+AI reporting guideline.", TRIPOD_URL)
    add_reference_paragraph(doc, 5, "PROBAST+AI risk-of-bias and applicability tool.", PROBAST_URL)
    add_reference_paragraph(doc, 6, "Albu et al. Dynamic prediction of CLABSI within seven days (BMC Infectious Diseases, 2025).", ALBU_URL)
    add_reference_paragraph(doc, 7, "Rahmani et al. Machine-learning prediction of CLABSI at 48 hours (AJIC, 2022).", RAHMANI_URL)

    path = OUT / "v0_5_run28_model_card.docx"
    doc.save(path)
    return path


def build_manuscript(e: dict, figures: dict[str, Path]) -> Path:
    doc = Document()
    configure_document(doc, "narrative_proposal")
    add_title_block(
        doc,
        "Temporal Prediction Of Seven-Day CVC-Associated BSI Proxy Risk In MIMIC-IV",
        "Manuscript-style results package for the frozen CVCML v0.5 model",
        f"Run 28 | Draft results and reporting framework | {date.today().isoformat()}",
    )
    p = doc.add_paragraph("Working status: evidence-consolidation draft. This document does not claim adjudicated NHSN CLABSI prediction or clinical readiness.")
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_heading("Structured Abstract", level=1)
    doc.add_heading("Background", level=2)
    doc.add_paragraph("Central-line-associated bloodstream infection prediction is difficult because events are uncommon, exposure changes over time, and retrospective electronic health record labels incompletely reproduce surveillance adjudication. We evaluated whether a temporally honest seven-day risk model could prioritize catheter episodes for infection-prevention review.")
    doc.add_heading("Methods", level=2)
    doc.add_paragraph("Recorded central venous catheter procedures in MIMIC-IV v3.1 were reconstructed into continuous exposure periods while retaining multiple episodes per ICU stay. Eligible episodes required at least 48 hours of recorded exposure. Daily landmarks were labeled for a strict primary-or-uncertain CVC-associated BSI proxy in the subsequent seven days. An XGBoost model using static context, laboratory, vital-sign, therapy, and line-timing features was calibrated using Platt scaling. Development used 2008-2013 data, calibration used 2014-2016, validation used 2017-2019, and the frozen model was evaluated once in a 2020-2022 temporal lockbox.")
    doc.add_heading("Results", level=2)
    lock = e["primary_lock"]
    doc.add_paragraph(f"The lockbox contained {int(lock.rows):,} eligible landmark rows from {int(lock.episodes):,} episodes and {int(lock.patients):,} patients, including {int(lock.positive_rows)} positive rows across {int(lock.positive_episodes)} positive episodes ({percent(lock.prevalence, 2)} row prevalence). The frozen model achieved ROC-AUC {metric(lock.roc_auc)}, PR-AUC {metric(lock.pr_auc)}, and {metric(lock.pr_auc_lift_over_prevalence, 2)}-fold PR-AUC lift over prevalence. Platt-calibrated Brier Skill Score was {metric(lock.brier_skill_score)}, calibration slope was {metric(lock.calibration_slope)}, and expected:observed ratio was {metric(lock.expected_observed_ratio)}. Reviewing the top 150 episodes produced PPV {percent(e['ep150'].precision_ppv)} and captured {int(e['ep150'].positive_episodes_captured)}/{int(e['ep150'].positive_episodes_total)} positive episodes ({percent(e['ep150'].episode_recall_sensitivity)}).")
    doc.add_heading("Conclusions", level=2)
    doc.add_paragraph("The model showed modest but temporally persistent risk stratification and clinically interpretable prioritization lift. Its current use case is a bounded infection-prevention review list rather than an interruptive bedside alarm. Outcome validity, external transportability, subgroup performance, and confidence intervals remain necessary before stronger clinical claims.")

    doc.add_heading("Study Design", level=1)
    add_table(doc, ["Component", "Specification"], [
        ["Database", "MIMIC-IV v3.1"],
        ["Exposure unit", "Continuous catheter-exposure episode reconstructed from recorded CVC procedures"],
        ["Eligibility", ">=48 hours of recorded catheter exposure"],
        ["Prediction time", "Daily landmark while eligible and at risk"],
        ["Outcome horizon", "Seven days"],
        ["Primary outcome", "Strict primary-or-uncertain CVC-associated BSI proxy"],
        ["Competing endpoints", "Landmarks censored at recorded line end, discharge, death, or end of observation as encoded by the v0.5 frame"],
        ["Model", "XGBoost; Platt calibration"],
        ["Temporal design", "Train 2008-2013; calibrate 2014-2016; validate 2017-2019; locked test 2020-2022"],
    ], [2600, 6760], header_fill=LIGHT_GRAY)

    doc.add_heading("Cohort Construction", level=1)
    a = e["episode_audit"]
    doc.add_paragraph(f"The episode redesign began with {int(a.raw_cvc_procedure_events):,} recorded CVC procedure events and produced {int(a.continuous_exposure_periods):,} continuous exposure periods. Of these, {int(a.eligible_48h_exposure_periods):,} met the 48-hour eligibility criterion. Retaining all reconstructed periods avoided selecting the eventual longest line, a future-information design that could discard episodes and misattribute cultures.")
    add_table(doc, ["Cohort stage", "Count", "Rate / note"], [
        ["Recorded CVC procedure events", f"{int(a.raw_cvc_procedure_events):,}", "Input records"],
        ["Continuous exposure periods", f"{int(a.continuous_exposure_periods):,}", "4-hour gap rule"],
        ["Eligible exposure periods", f"{int(a.eligible_48h_exposure_periods):,}", ">=48 hours"],
        ["Broad proxy-positive episodes", f"{int(a.broad_proxy_positive_episodes):,}", percent(a.broad_proxy_positive_rate, 2)],
        ["Strict organism-rule positives", f"{int(a.strict_proxy_positive_episodes):,}", percent(a.strict_proxy_positive_rate, 2)],
        ["Primary-or-uncertain source-screened positives", "291", "1.28% of all reconstructed periods"],
    ], [4100, 1800, 3460], header_fill=LIGHT_GRAY)

    doc.add_heading("Outcome Definition", level=1)
    doc.add_paragraph("The primary outcome was intentionally named a strict CVC-associated BSI proxy. It combined catheter-exposure timing, blood-culture timing, organism rules for recognized pathogens and repeated common commensals, and structured secondary-source screening. Source screening separated primary-likely, secondary-possible, and uncertain episodes; the primary model combined primary-likely and uncertain events to avoid treating incomplete structured evidence as definitive exclusion.")
    add_callout(doc, "Why this is not adjudicated CLABSI", "NHSN CLABSI determination additionally depends on eligible-line rules, primary versus secondary BSI attribution, symptom criteria for common commensals, MBI-LCBI handling, and surveillance review. MIMIC-IV does not contain an official infection-prevention adjudication field. The project therefore preserves proxy terminology throughout.", PALE_RED)

    doc.add_heading("Model And Predictors", level=1)
    doc.add_paragraph("The frozen model used 173 pre-encoding candidate fields: 168 numeric and five categorical inputs. Feature families included demographics and admission context; catheter timing and type; rolling laboratory summaries; vital-sign summaries; measurement recency and counts; and therapy context such as antibiotics, vasopressors, ventilation, and fluid-related variables. XGBoost hyperparameters were fixed during development, and Platt scaling was fitted only on the 2014-2016 calibration period.")
    add_table(doc, ["Parameter", "Value"], [
        ["n_estimators", "300"], ["max_depth", "3"], ["learning_rate", "0.03"], ["min_child_weight", "10"],
        ["subsample", "0.80"], ["colsample_bytree", "0.80"], ["reg_alpha", "0.5"], ["reg_lambda", "1.0"],
        ["class handling", "scale_pos_weight from training negatives/positives"], ["tree method", "hist"],
    ], [3000, 6360], header_fill=LIGHT_GRAY)
    top = e["top_importance"]
    add_table(doc, ["Rank", "Feature", "Gain importance"], [[str(j + 1), clean_feature(r.feature), metric(r.importance, 4)] for j, (_, r) in enumerate(top.head(15).iterrows())], [900, 6500, 1960], header_fill=LIGHT_GRAY)
    doc.add_paragraph("Feature importance is descriptive rather than causal. Measurement-count and therapy-context predictors may encode severity and clinician behavior, which can improve ranking while reducing transportability. The outcome-adjacent early_positive_culture field requires a dedicated timestamp and ablation audit.")

    doc.add_heading("Temporal Evaluation", level=1)
    val = e["primary_val"]
    add_table(doc, ["Measure", "Validation 2017-2019", "Lockbox 2020-2022"], [
        ["Rows", f"{int(val.rows):,}", f"{int(lock.rows):,}"],
        ["Episodes", "648 total; 54 positive", f"{int(lock.episodes):,} total; {int(lock.positive_episodes)} positive"],
        ["Prevalence", percent(val.prevalence, 2), percent(lock.prevalence, 2)],
        ["ROC-AUC", metric(val.roc_auc), metric(lock.roc_auc)],
        ["PR-AUC", metric(val.pr_auc), metric(lock.pr_auc)],
        ["PR-AUC lift", f"{metric(val.pr_auc_lift_over_prevalence, 2)}x", f"{metric(lock.pr_auc_lift_over_prevalence, 2)}x"],
        ["Brier Skill Score", metric(val.brier_skill_score), metric(lock.brier_skill_score)],
        ["Expected:observed", metric(val.expected_observed_ratio), metric(lock.expected_observed_ratio)],
        ["Calibration slope", "Not reported due unstable estimate", metric(lock.calibration_slope)],
    ], [3300, 3030, 3030], header_fill=LIGHT_GRAY)
    add_figure(doc, figures["validation_lockbox"], "Figure 1. Validation and frozen temporal-lockbox performance. PR-AUC should be interpreted relative to each split's prevalence.")
    add_figure(doc, figures["calibration"], "Figure 2. Lockbox calibration by risk decile after Platt scaling.")
    doc.add_paragraph("The lockbox ROC-AUC remained similar to validation, while PR-AUC increased from 0.069 to 0.110 alongside an increase in prevalence from 4.25% to 5.92%. The resulting 1.85-fold PR-AUC lift indicates nontrivial ranking beyond baseline prevalence, but absolute discrimination remains modest. The positive lockbox sample comprised only 26 episodes, so uncertainty intervals are still required.")

    doc.add_heading("Operational Evaluation", level=1)
    add_table(doc, ["Policy", "Reviews", "PPV", "Episode recall", "False reviews / TP"], [
        ["Top 5% landmark rows", f"{int(e['row5'].rows_reviewed)} rows", percent(e['row5'].precision_ppv), percent(e['row5'].episode_recall_sensitivity), metric(e['row5'].false_reviews_per_true_positive, 1)],
        ["Top 10% landmark rows", f"{int(e['row10'].rows_reviewed)} rows", percent(e['row10'].precision_ppv), percent(e['row10'].episode_recall_sensitivity), metric(e['row10'].false_reviews_per_true_positive, 1)],
        ["Top 100 episodes", "100 episodes", percent(e['ep100'].precision_ppv), percent(e['ep100'].episode_recall_sensitivity), metric(e['ep100'].false_reviews_per_true_positive, 1)],
        ["Top 150 episodes", "150 episodes", percent(e['ep150'].precision_ppv), percent(e['ep150'].episode_recall_sensitivity), metric(e['ep150'].false_reviews_per_true_positive, 1)],
        ["Top 250 episodes", "250 episodes", percent(e['ep250'].precision_ppv), percent(e['ep250'].episode_recall_sensitivity), metric(e['ep250'].false_reviews_per_true_positive, 1)],
    ], [2800, 1650, 1350, 1700, 1860], header_fill=LIGHT_GRAY)
    doc.add_paragraph("The top-episode policies are the more coherent operational framing because they bound daily review burden and avoid repeated alerts from the same episode. The top-150 policy reviewed 55.6% of all lockbox episodes, captured 57.7% of positive episodes, and had 10.0% PPV. This may support retrospective or infection-prevention triage, but it is too burdensome and imprecise for an interruptive bedside alert.")

    doc.add_heading("Locked Error Analysis", level=1)
    doc.add_paragraph("At the top-10% row policy, 31 positive rows were reviewed, 196 reviewed rows were false positives, and 103 positive rows were missed. False positives frequently reflected high measurement intensity and treatment context rather than random noise. Missed positives showed greater source uncertainty and, in several features, less intense measurement. These patterns suggest the model partly learns clinician attention and severity, while some bloodstream infections remain difficult to anticipate from structured pre-event data.")
    add_figure(doc, figures["error_contrast"], "Figure 3. Standardized feature contrasts between reviewed negatives and missed positives in the locked test set.")

    doc.add_heading("Proxy-Label Validation", level=1)
    icd, ge, gh = e["icd_primary"], e["grain_episode"], e["grain_hadm"]
    add_table(doc, ["Grain", "Units", "Proxy +", "ICD +", "Both +", "Jaccard"], [
        ["Catheter episode", f"{int(ge.n_units):,}", f"{int(ge.proxy_positive):,}", f"{int(ge.icd_positive):,}", f"{int(ge.both_positive):,}", percent(ge.positive_set_jaccard, 1)],
        ["Hospital admission", f"{int(gh.n_units):,}", f"{int(gh.proxy_positive):,}", f"{int(gh.icd_positive):,}", f"{int(gh.both_positive):,}", percent(gh.positive_set_jaccard, 1)],
    ], [2300, 1500, 1400, 1400, 1300, 1460], header_fill=LIGHT_GRAY)
    add_figure(doc, figures["icd_overlap"], "Figure 4. Episode-level proxy and ICD-code overlap.")
    add_figure(doc, figures["icd_era"], "Figure 5. Discordance by era. ICD-positive/proxy-negative discrepancies decreased in later years, emphasizing temporal coding drift.")
    doc.add_paragraph("Only 60 catheter episodes were positive under both the primary proxy and the specific ICD comparator. Positive-set Jaccard agreement was 6.6% at the episode level and 8.5% after reconciling both labels to hospital admission. The 167 ICD-positive admissions containing multiple catheter episodes created unavoidable attribution ambiguity. Because MIMIC diagnoses are assigned at hospital-admission grain after review of the completed stay, ICD codes were treated as a secondary agreement check rather than ground truth.")

    doc.add_heading("Comparison With Prior Work", level=1)
    add_table(doc, ["Study", "Task / evaluation", "Reported result", "Comparability caveat"], [
        ["Rahmani et al., 2022", "CLABSI risk at 48 h; 27,619 encounters; 80/20 split", "XGBoost AUROC 0.762", "Different horizon, label, sites, and non-temporal split"],
        ["Albu et al., 2025", "7-day dynamic CLABSI prediction; temporal evaluation", "Best individual AUROC 0.748; superlearner 0.751", "Different institution, adjudication, competing-risk design, and prevalence"],
        ["CVCML v0.5", "7-day strict CVC-associated BSI proxy; 2020-2022 temporal lockbox", "AUROC 0.623; PR-AUC 0.110; lift 1.85x", "MIMIC-IV proxy label and 26 positive lockbox episodes"],
    ], [2100, 3000, 2100, 2160], header_fill=LIGHT_GRAY, font_size=8)
    doc.add_paragraph("These values are contextual benchmarks, not head-to-head comparisons. Outcome definitions, cohort ascertainment, prediction times, institution, and validation design differ materially. The lower CVCML AUROC should be interpreted alongside its later-time lockbox, transparent proxy limitations, and explicit operational burden analysis.")

    doc.add_heading("Strengths", level=1)
    add_table(doc, ["Strength", "Why it matters"], [
        ["All reconstructed exposure periods retained", "Avoids future-informed longest-line selection and reduces episode misattribution."],
        ["Daily seven-day landmark design", "Aligns predictions with a clinically actionable forward window."],
        ["Temporal lockbox", "Separates model development from later-time evaluation."],
        ["Calibration and Brier Skill Score", "Evaluates probability quality against a prevalence-only reference."],
        ["Review-burden analysis", "Translates statistical ranking into workload, PPV, and episode capture."],
        ["Label sensitivity and ICD agreement", "Makes outcome uncertainty visible instead of presenting a proxy as adjudicated truth."],
    ], [3300, 6060], header_fill=LIGHT_GRAY)

    doc.add_heading("Limitations", level=1)
    limits = read_csv(OUT / "v0_5_run28_limitations_and_actions.csv")
    add_table(doc, ["Domain", "Limitation", "Action before stronger claim"], [[r.domain, r.current_limitation, r.required_action] for _, r in limits.iterrows()], [1800, 3800, 3760], header_fill=LIGHT_GRAY, font_size=8)

    doc.add_heading("Recommended Next Decision", level=1)
    doc.add_paragraph("Run 28 completes the agreed evidence-consolidation pathway. The model is now sufficiently characterized to support a portfolio case study and a manuscript methods/results draft, provided that the outcome is consistently called a strict CVC-associated BSI proxy and the work is presented as risk prioritization rather than clinical diagnosis.")
    add_callout(doc, "Recommended focused improvement", "Before pursuing a stronger clinical manuscript claim, perform one bounded validity run: audit and exclude early_positive_culture, then adjudicate a stratified sample of proxy/ICD concordant and discordant cases using notes and structured source evidence. Avoid reopening broad feature tuning against the temporal lockbox.", PALE_GREEN)
    doc.add_paragraph("If external transportability is the priority, the next substantive model experiment should be validation of the frozen pipeline on a second database. Further MIMIC-only tuning would add less evidence than an external test.")

    doc.add_heading("References", level=1)
    add_reference_paragraph(doc, 1, "CDC/NHSN Patient Safety Component Manual, Bloodstream Infection Event chapter (January 2026).", CDC_URL)
    add_reference_paragraph(doc, 2, "MIMIC-IV diagnoses_icd documentation.", MIMIC_DIAG_URL)
    add_reference_paragraph(doc, 3, "MIMIC-IV procedureevents documentation.", MIMIC_PROC_URL)
    add_reference_paragraph(doc, 4, "Collins et al. TRIPOD+AI reporting guideline. BMJ 2024.", TRIPOD_URL)
    add_reference_paragraph(doc, 5, "Moons et al. PROBAST+AI. BMJ 2025.", PROBAST_URL)
    add_reference_paragraph(doc, 6, "Albu et al. Dynamic prediction of CLABSI within seven days. BMC Infectious Diseases 2025.", ALBU_URL)
    add_reference_paragraph(doc, 7, "Rahmani et al. Prediction of CLABSI using EHR data. American Journal of Infection Control 2022.", RAHMANI_URL)

    path = OUT / "v0_5_run28_manuscript_results_package.docx"
    doc.save(path)
    return path


def write_markdown(e: dict) -> list[Path]:
    lock = e["primary_lock"]
    model_card = OUT / "v0_5_run28_model_card.md"
    model_card_text = (
        "# CVCML v0.5 Model Card\n\n"
        "**Status:** Research prototype; not for clinical use.\n\n"
        "## Intended use\n\nBounded infection-prevention review-list prioritization for seven-day strict CVC-associated BSI proxy risk. Not an autonomous diagnosis or bedside alarm.\n\n"
        "## Frozen model\n\n- XGBoost with Platt calibration\n- Daily landmark prediction\n- Development: 2008-2013; calibration: 2014-2016; validation: 2017-2019\n- Locked temporal test: 2020-2022\n- Primary target: future_strict_primary_or_uncertain_cvc_bsi_proxy_7d\n\n"
        f"## Lockbox performance\n\n- Rows: {int(lock.rows):,}; episodes: {int(lock.episodes):,}; positive episodes: {int(lock.positive_episodes)}\n"
        f"- ROC-AUC: {metric(lock.roc_auc)}\n- PR-AUC: {metric(lock.pr_auc)} ({metric(lock.pr_auc_lift_over_prevalence, 2)}x prevalence)\n"
        f"- Brier Skill Score: {metric(lock.brier_skill_score)}\n- Calibration intercept/slope: {metric(lock.calibration_intercept)} / {metric(lock.calibration_slope)}\n"
        f"- Expected:observed: {metric(lock.expected_observed_ratio)}\n\n"
        "## Key limitations\n\n- Proxy outcome, not adjudicated NHSN CLABSI.\n- procedureevents does not provide a complete line denominator.\n"
        "- early_positive_culture is outcome-adjacent and requires an exclusion audit.\n- Single-institution temporal validation and only 26 positive lockbox episodes.\n"
        "- Final cluster-bootstrap confidence intervals and subgroup evaluation remain outstanding.\n\n"
        "## Decision\n\nCredible for a portfolio/manuscript draft and retrospective prioritization research. Not ready for clinical deployment.\n"
    )
    model_card.write_text(model_card_text, encoding="utf-8")
    manuscript = OUT / "v0_5_run28_manuscript_results.md"
    manuscript_text = (
        "# Manuscript Results Summary\n\n"
        "The frozen v0.5 model was evaluated on 2,262 eligible daily landmark rows from 270 catheter episodes in 2020-2022. The lockbox included 134 positive rows across 26 positive episodes (5.92% row prevalence). ROC-AUC was 0.623 and PR-AUC was 0.110, corresponding to 1.85-fold lift over prevalence. Platt-calibrated Brier Skill Score was 0.013, calibration slope was 1.056, and expected:observed ratio was 0.888.\n\n"
        "At a target-aligned top-150 episode review budget, PPV was 10.0% and 15 of 26 positive episodes were captured (57.7%), with 9.0 false reviews per true-positive row. The appropriate use case is a bounded infection-prevention review list rather than an interruptive bedside alarm.\n\n"
        "ICD-code agreement was low (episode-level positive-set Jaccard 6.6%; admission-level 8.5%), but ICD diagnoses are admission-level discharge codes and were not treated as ground truth. The outcome remains a strict CVC-associated BSI proxy.\n\n"
        "The next focused validity step is notes/source adjudication plus exclusion of the outcome-adjacent early_positive_culture feature. External validation would add more evidence than additional MIMIC-only tuning.\n"
    )
    manuscript.write_text(manuscript_text, encoding="utf-8")
    return [model_card, manuscript]


def qa_docx(path: Path) -> dict:
    doc = Document(path)
    empty_tables = []
    geometry_issues = []
    for i, table in enumerate(doc.tables):
        if not table.rows or not table.columns:
            empty_tables.append(i)
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        if tbl_w is None or tbl_w.get(qn("w:w")) != "9360":
            geometry_issues.append(i)
    searchable_text = [p.text for p in doc.paragraphs]
    searchable_text.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    searchable_text.extend(p.text for section in doc.sections for p in section.footer.paragraphs)
    return {
        "path": str(path),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "sections": len(doc.sections),
        "empty_tables": empty_tables,
        "table_geometry_issues": geometry_issues,
        "has_research_warning": any("not for clinical use" in text.lower() for text in searchable_text),
    }


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    evidence = load_evidence()
    figures = copy_figures()
    table_paths = write_consolidated_tables(evidence)
    model_card = build_model_card(evidence, figures, table_paths)
    manuscript = build_manuscript(evidence, figures)
    markdown_paths = write_markdown(evidence)

    qa = [qa_docx(model_card), qa_docx(manuscript)]
    qa_path = OUT / "v0_5_run28_docx_structural_qa.json"
    qa_path.write_text(json.dumps(qa, indent=2), encoding="utf-8")
    if any(item["empty_tables"] or item["table_geometry_issues"] or not item["has_research_warning"] for item in qa):
        raise RuntimeError(f"DOCX structural QA failed: {qa}")

    manifest = {
        "run": "Run 28",
        "scope": "Model card and manuscript-style results package for the frozen CVCML v0.5 model",
        "lockbox_status": "No refit, recalibration, threshold tuning, or model selection performed",
        "documents": [str(model_card), str(manuscript)],
        "markdown": [str(p) for p in markdown_paths],
        "tables": {k: str(v) for k, v in table_paths.items()},
        "figures": {k: str(v) for k, v in figures.items()},
        "qa": str(qa_path),
        "primary_result": {
            "roc_auc": float(evidence["primary_lock"].roc_auc),
            "pr_auc": float(evidence["primary_lock"].pr_auc),
            "pr_auc_lift": float(evidence["primary_lock"].pr_auc_lift_over_prevalence),
            "brier_skill_score": float(evidence["primary_lock"].brier_skill_score),
            "calibration_slope": float(evidence["primary_lock"].calibration_slope),
            "expected_observed_ratio": float(evidence["primary_lock"].expected_observed_ratio),
        },
    }
    (OUT / "v0_5_run28_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    notes = OUT / "run_28_v0_5_model_card_manuscript_notes.md"
    notes.write_text(
        "# Run 28 - v0.5 Model Card and Manuscript Package\n\n"
        "Run 28 is an evidence-consolidation run. It did not refit, recalibrate, or tune the frozen model and did not reuse the temporal lockbox for development.\n\n"
        "Primary locked result: ROC-AUC 0.623, PR-AUC 0.110 (1.85x prevalence), Brier Skill Score 0.013, calibration slope 1.056, and E:O 0.888 in 2,262 rows from 270 episodes with 26 positive episodes.\n\n"
        "Primary operating results use the target-aligned Run 25 estimand. Higher episode-capture values from Run 26 are clearly marked exploratory because they use a broader post-hoc episode estimand.\n\n"
        "The package consistently names the outcome a strict CVC-associated BSI proxy and treats ICD agreement as convergent validity, not ground truth.\n",
        encoding="utf-8",
    )
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()

