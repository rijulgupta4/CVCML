from __future__ import annotations

import csv
import hashlib
import json
import shutil
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r"C:\path\to\CVCML")
WORKSPACE = Path(r"C:\path\to\workspace")
STAGE = WORKSPACE / "Run 33 (v0.5 Publication Consolidation)"
FIGURES = STAGE / "figures"

RUN28 = PROJECT / "Outputs" / "Run 28 (v0.5 Model Card and Manuscript Package)"
RUN29 = PROJECT / "Outputs" / "Run 29 (v0.5 Outcome Validity and Leakage Audit)"
RUN30 = PROJECT / "Outputs" / "Run 30 (v0.5 Safe Candidate Characterization)"
RUN31 = PROJECT / "Outputs" / "Run 31 (External Validation Feasibility)"
RUN32 = PROJECT / "Outputs" / "Run 32 (External Label Transportability)"
HUB_SOURCE = WORKSPACE / "CVCML Project Hub - updated v0.5 Run32.docx"
HUB_OUTPUT = WORKSPACE / "CVCML Project Hub - updated v0.5 Run33.docx"

BLUE = "1F4E79"
MID_BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
VERY_LIGHT = "F4F6F9"
INK = "233240"
MUTED = "5B6573"
GREEN = "2F6B4F"
GOLD = "806000"
RED = "9B1C1C"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_dxa = sum(round(width * 1440) for width in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def configure_document(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Title", 24, BLUE, 0, 8),
        ("Subtitle", 11, MUTED, 0, 12),
        ("Heading 1", 16, MID_BLUE, 15, 7),
        ("Heading 2", 12.5, BLUE, 11, 5),
        ("Heading 3", 11, BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = short_title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("CVCML research prototype | Run 33 | ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    doc.add_paragraph(title, style="Title")
    doc.add_paragraph(subtitle, style="Subtitle")


def add_callout(doc: Document, label: str, text: str, fill=VERY_LIGHT, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_geometry(table, [6.8])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margin(cell, 120, 150, 120, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(f"{label}: ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string(accent)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str, bold_lead: str | None = None) -> None:
    try:
        p = doc.add_paragraph(style="List Bullet")
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run("- ")
    p.paragraph_format.space_after = Pt(4)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    try:
        p = doc.add_paragraph(style="List Number")
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run("- ")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float], font_size=8.7) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        shade_cell(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margin(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(BLUE)
            run.font.size = Pt(font_size)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for values in rows:
        row = table.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, path: Path, caption: str, width=6.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    for run in cap.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(MUTED)


def copy_figures() -> dict[str, Path]:
    FIGURES.mkdir(parents=True, exist_ok=True)
    sources = {
        "result_availability": RUN29 / "plots" / "v0_5_run29_result_unavailability_by_landmark.png",
        "calibration": RUN30 / "plots" / "v0_5_run30_calibration_deciles.png",
        "review_policy": RUN30 / "plots" / "v0_5_run30_episode_review_policy.png",
        "subgroups": RUN30 / "plots" / "v0_5_run30_subgroup_pr_auc_lift.png",
        "organism_profile": RUN32 / "plots" / "run32_mimic_armd_organism_profile.png",
        "source_window": RUN32 / "plots" / "run32_secondary_source_window_sensitivity.png",
    }
    copied = {}
    for key, source in sources.items():
        destination = FIGURES / source.name
        shutil.copy2(source, destination)
        copied[key] = destination
    return copied


def load_data():
    overall = pd.read_csv(RUN30 / "v0_5_run30_overall_metric_ci.csv")
    policy = pd.read_csv(RUN30 / "v0_5_run30_episode_review_policy_ci.csv")
    ablation = pd.read_csv(RUN29 / "v0_5_run29_ablation_model_comparison.csv")
    windows = pd.read_csv(RUN32 / "run32_secondary_source_window_sensitivity.csv")
    external_scope = pd.read_csv(RUN32 / "run32_external_validation_scope.csv")
    with (RUN32 / "run32_metadata.json").open(encoding="utf-8") as handle:
        external_meta = json.load(handle)
    cohort = pd.read_csv(RUN28 / "v0_5_run28_cohort_summary.csv")
    return overall, policy, ablation, windows, external_scope, external_meta, cohort


def metric_row(overall: pd.DataFrame, metric: str) -> pd.Series:
    return overall.loc[overall.metric == metric].iloc[0]


def policy_value(policy: pd.DataFrame, top_percent: int, metric: str) -> pd.Series:
    return policy.loc[(policy.top_percent == top_percent) & (policy.metric == metric)].iloc[0]


def write_csvs(overall, policy, ablation, windows, external_meta, cohort) -> None:
    performance = overall[["metric", "estimate", "ci_lower_95", "ci_upper_95", "landmark_rows", "subjects", "episodes"]].copy()
    performance.insert(0, "evaluation_period", "2017-2019 development-validation")
    performance.insert(1, "model_variant", "safe_exclude_early_positive")
    performance.to_csv(STAGE / "v0_5_run33_performance_summary.csv", index=False)

    policy.to_csv(STAGE / "v0_5_run33_review_policy_summary.csv", index=False)

    external_rows = [
        ["ARMD-MGB", "External label-component transportability", "Positive blood-culture accession", int(external_meta["armd_positive_blood_accessions"]), "Organism rule and urine/respiratory source-screen sensitivity", "Supported"],
        ["ARMD-MGB", "External full-model validation", "Catheter episode", "NA", "No catheter exposure or predictor timeline", "Not supported"],
        ["eICU-CRD", "External full-model validation feasibility", "ICU stay", 1, "Only one positive-blood stay met >=48 h explicit-placement criterion", "Failed feasibility gate"],
    ]
    pd.DataFrame(external_rows, columns=["database", "assessment", "unit", "relevant_n", "evidence", "conclusion"]).to_csv(
        STAGE / "v0_5_run33_external_validation_summary.csv", index=False
    )

    claims = [
        ["Leakage audit", "early_positive_culture was outcome/result-derived and was excluded from the manuscript-safe model.", "Run 29", "Supported", "Development-only audit; storetime is last known update, not first clinical notification."],
        ["Internal discrimination", "The leakage-safe candidate showed modest discrimination in 2017-2019.", "Run 30", "Supported", "ROC-AUC 0.612; PR-AUC 0.065; wide patient-clustered CIs."],
        ["Operational use", "The score may enrich a bounded infection-prevention review list.", "Run 30", "Supported with caution", "Top 10% PPV 15.4%, recall 18.5%; not suitable for interruptive alarms."],
        ["Absolute risk", "The score is a reliable absolute risk estimator.", "Run 30", "Not supported", "Brier Skill Score near zero; calibration estimates are imprecise."],
        ["Outcome validity", "The target is adjudicated NHSN CLABSI.", "Runs 22, 29, 32", "Not supported", "Use strict CVC-associated BSI proxy terminology."],
        ["External label transportability", "Organism and partial source-screen logic can be implemented in ARMD-MGB.", "Run 32", "Supported", "Different unit and case mix; not catheter-episode prevalence."],
        ["External model validation", "The frozen model has been validated at another institution.", "Runs 31-32", "Not supported", "eICU failed feasibility; ARMD-MGB lacks predictor and catheter timelines."],
        ["Publication readiness", "The project is ready for a transparent retrospective methods/results manuscript draft.", "Run 33", "Supported", "Requires external full-model validation or explicit limitation before stronger clinical claims."],
    ]
    pd.DataFrame(claims, columns=["claim_area", "claim", "evidence", "status", "boundary"]).to_csv(
        STAGE / "v0_5_run33_claim_register.csv", index=False
    )

    evidence = [
        ["Catheter exposure reconstruction", "Run 16", "v0_5_episode_label_audit.csv", "All recorded catheter exposure periods; no longest-line selection"],
        ["Daily landmark frame", "Run 17", "v0_5_daily_landmark_temporal_split_audit.csv", "One seven-day target with daily landmarks"],
        ["Secondary-source proxy logic", "Run 22", "v0_5_secondary_source_label_audit.csv", "Broad, strict primary, secondary-possible, and uncertain strata"],
        ["Frozen development model", "Run 23", "v0_5_run23_label_sensitivity_model_comparison.csv", "Model specification inherited without Run 29 retuning"],
        ["Historical temporal sensitivity", "Run 25", "v0_5_run25_lockbox_model_comparison.csv", "Previously inspected 2020-2022 period; not a pristine safe-model lockbox"],
        ["Leakage audit", "Run 29", "v0_5_run29_ablation_model_comparison.csv", "Feature exclusion and result-availability audit"],
        ["Safe-candidate uncertainty", "Run 30", "v0_5_run30_overall_metric_ci.csv", "2,000 patient-clustered bootstrap replicates"],
        ["External feasibility", "Run 31", "run31b_eicu_feasibility_summary.csv", "eICU exact-validation gate"],
        ["External label transportability", "Run 32", "run32_mimic_armd_organism_comparison.csv", "ARMD-MGB organism distribution and source-screen sensitivity"],
        ["Publication consolidation", "Run 33", "v0_5_run33_claim_register.csv", "Current claim boundaries and manuscript-safe language"],
    ]
    pd.DataFrame(evidence, columns=["claim_area", "run", "source_file", "role"]).to_csv(
        STAGE / "v0_5_run33_evidence_traceability.csv", index=False
    )

    limitations = [
        ["Outcome", "Strict CVC-associated BSI proxy, not adjudicated NHSN CLABSI", "High", "Retain proxy terminology; perform blinded chart/notes adjudication where possible"],
        ["Line denominator", "procedureevents presence is positive evidence; absence does not prove no line", "High", "Validate episode ascertainment against a richer EHR or local registry"],
        ["External performance", "No suitable external dataset reproduced catheter episodes plus predictors and outcomes", "High", "Seek longitudinal hospital EHR with line timing and microbiology"],
        ["Temporal test", "The 2020-2022 period was inspected repeatedly before leakage-safe finalization", "High", "Treat as post-hoc sensitivity; reserve a truly untouched cohort for future validation"],
        ["Discrimination", "Safe candidate has modest and imprecise discrimination", "Moderate", "Avoid further broad MIMIC tuning; validate transportability and failure modes"],
        ["Calibration", "Brier Skill Score is approximately zero and calibration CIs are wide", "High", "Do not present probabilities as bedside absolute risk without external recalibration"],
        ["Review burden", "Most reviewed episodes are false positives", "Moderate", "Use bounded infection-prevention review queues, not interruptive alerts"],
        ["Source attribution", "Urine/respiratory culture matches do not prove secondary BSI", "High", "Add notes, wound/abdominal sources, symptom rules, and MBI-LCBI review"],
    ]
    pd.DataFrame(limitations, columns=["domain", "limitation", "priority", "recommended_action"]).to_csv(
        STAGE / "v0_5_run33_limitations_and_actions.csv", index=False
    )

    cohort.to_csv(STAGE / "v0_5_run33_cohort_summary.csv", index=False)
    windows.to_csv(STAGE / "v0_5_run33_external_source_window_summary.csv", index=False)


def build_model_card(overall, policy, external_meta, figures) -> Path:
    doc = Document()
    configure_document(doc, "CVCML v0.5 Leakage-Safe Model Card")
    add_title(doc, "CVCML v0.5 Leakage-Safe Model Card", "Run 33 publication consolidation | Updated 17 July 2026")
    add_callout(doc, "Status", "Research prototype for retrospective evaluation. Not for clinical use, autonomous diagnosis, or interruptive bedside alerting.", fill="FCE8E6", accent=RED)

    doc.add_heading("Intended Use", level=1)
    doc.add_paragraph(
        "Prioritize a bounded daily infection-prevention review list for catheter episodes with elevated seven-day risk of the study's strict primary-or-uncertain CVC-associated bloodstream infection proxy. The score is best interpreted as a ranking signal for review capacity, not as a diagnosis or calibrated bedside probability."
    )
    add_bullet(doc, "Primary user: infection-prevention or clinical analytics teams conducting retrospective or silent prospective review.")
    add_bullet(doc, "Permitted action: rank episodes for review under a prespecified daily or weekly capacity.")
    add_bullet(doc, "Prohibited action: trigger treatment, line removal, or a nurse-facing interruptive alert without clinical assessment and prospective validation.")

    doc.add_heading("Model and Prediction Task", level=1)
    add_table(doc, ["Element", "Specification"], [
        ["Population", "MIMIC-IV ICU admissions with recorded central-line procedure events reconstructed into continuous exposure periods"],
        ["Unit", "Daily landmark nested within an eligible catheter exposure episode"],
        ["Horizon", "Strict primary-or-uncertain CVC-associated BSI proxy in the next 7 days"],
        ["Model", "Frozen XGBoost candidate with Platt calibration; Run 29 safe variant excludes early_positive_culture"],
        ["Development", "2008-2013 training; 2014-2016 calibration"],
        ["Current evaluation", "2017-2019 development-validation; 2,000 patient-clustered bootstrap replicates"],
        ["Historical sensitivity", "2020-2022 was repeatedly inspected before the safe model was finalized and is not a pristine lockbox"],
    ], [1.55, 5.25], font_size=8.5)

    doc.add_heading("Outcome Definition", level=1)
    doc.add_paragraph(
        "The outcome is a strict CVC-associated BSI proxy: an eligible reconstructed line-exposure episode, positive blood culture after at least 48 hours of exposure and before censoring, organism rules for recognized pathogens or repeated common commensals, and partial secondary-source screening. It is not NHSN-adjudicated CLABSI."
    )
    add_callout(doc, "Why the terminology matters", "NHSN determination also requires primary-versus-secondary attribution, symptom rules for common commensals, MBI-LCBI assessment, eligible-line rules, and surveillance adjudication that cannot be fully reconstructed from available structured MIMIC-IV fields.", fill="FFF4CE", accent=GOLD)

    doc.add_heading("Leakage Audit", level=1)
    doc.add_paragraph(
        "Run 29 found that early_positive_culture encoded eventual culture positivity at the episode level. Among development landmark rows carrying the flag, 42.4% did not yet have an organism-positive storetime by the landmark. Excluding the feature reduced validation PR-AUC by only 0.0038, so the leakage-safe variant became the manuscript candidate without retuning."
    )
    add_figure(doc, figures["result_availability"], "Figure 1. Availability audit for the outcome-adjacent culture feature across landmark times.", width=6.1)

    roc = metric_row(overall, "roc_auc")
    pr = metric_row(overall, "pr_auc")
    lift = metric_row(overall, "pr_auc_lift")
    bss = metric_row(overall, "brier_skill_score")
    slope = metric_row(overall, "calibration_slope")
    eo = metric_row(overall, "expected_observed_ratio")

    doc.add_heading("Performance", level=1)
    add_table(doc, ["Measure", "Estimate", "Patient-clustered 95% CI", "Interpretation"], [
        ["ROC-AUC", f"{roc.estimate:.3f}", f"{roc.ci_lower_95:.3f}-{roc.ci_upper_95:.3f}", "Modest discrimination; CI includes weak performance"],
        ["PR-AUC", f"{pr.estimate:.3f}", f"{pr.ci_lower_95:.3f}-{pr.ci_upper_95:.3f}", "Above 4.25% row prevalence"],
        ["PR-AUC lift", f"{lift.estimate:.2f}x", f"{lift.ci_lower_95:.2f}-{lift.ci_upper_95:.2f}x", "Ranking signal is present but limited"],
        ["Brier Skill Score", f"{bss.estimate:.3f}", f"{bss.ci_lower_95:.3f}-{bss.ci_upper_95:.3f}", "No clear improvement over constant-prevalence prediction"],
        ["Calibration slope", f"{slope.estimate:.3f}", f"{slope.ci_lower_95:.3f}-{slope.ci_upper_95:.3f}", "Very imprecise; absolute-risk use is unsupported"],
        ["Expected:observed", f"{eo.estimate:.3f}", f"{eo.ci_lower_95:.3f}-{eo.ci_upper_95:.3f}", "Average predicted events exceed observed events"],
    ], [1.45, 1.0, 1.75, 2.6], font_size=8.3)
    doc.add_paragraph(
        "Evaluation included 5,694 landmarks from 648 episodes and 590 patients; 242 landmarks were positive (4.25%), and 54 episodes were positive. Confidence intervals condition on the fitted model and calibration map and do not include model-development uncertainty."
    )
    add_figure(doc, figures["calibration"], "Figure 2. Calibration by risk group for the leakage-safe candidate in 2017-2019.", width=6.1)

    doc.add_heading("Review-List Policy", level=1)
    p10_ppv = policy_value(policy, 10, "episode_ppv")
    p10_rec = policy_value(policy, 10, "episode_recall")
    p10_false = policy_value(policy, 10, "false_reviews_per_true_positive")
    add_table(doc, ["Review budget", "Episodes", "PPV", "Positive-episode recall", "False reviews / TP"], [
        ["Top 5%", 33, "15.2%", "9.3%", "5.60"],
        ["Top 10%", 65, f"{p10_ppv.estimate:.1%}", f"{p10_rec.estimate:.1%}", f"{p10_false.estimate:.2f}"],
        ["Top 20%", 130, "16.2%", "38.9%", "5.19"],
    ], [1.25, 1.0, 1.15, 1.8, 1.6], font_size=8.5)
    doc.add_paragraph(
        "At the top 10% budget, PPV was 15.4% (95% CI 7.8%-26.9%) and recall was 18.5% (10.6%-31.5%). This may be useful for a bounded review queue, but approximately 5.5 false reviews occurred per true positive."
    )
    add_figure(doc, figures["review_policy"], "Figure 3. Episode-level review yield and burden across fixed review budgets.", width=6.1)

    doc.add_heading("External Evidence", level=1)
    add_bullet(doc, f"ARMD-MGB included {int(external_meta['armd_positive_blood_accessions']):,} positive blood-culture accessions; {int(external_meta['armd_strict_accessions']):,} met the finalized organism rule.")
    add_bullet(doc, f"Across 28 shared organism categories, rank correlation with MIMIC-IV was rho={float(external_meta['shared_organism_spearman_rho']):.3f}; Jensen-Shannon divergence was {float(external_meta['jensen_shannon_divergence']):.3f}, demonstrating both coherence and case-mix shift.")
    add_bullet(doc, "eICU could not reproduce the task: only one positive-blood stay met the >=48-hour explicit-line-placement feasibility criterion.")
    add_callout(doc, "Claim boundary", "The model has MIMIC-IV internal/temporal evidence and external validation of label components. It does not have external institutional validation of discrimination, calibration, or review burden.", fill="FFF4CE", accent=GOLD)

    doc.add_heading("Known Limitations", level=1)
    for text in [
        "Single-center MIMIC-IV development and no feasible full-model external validation cohort.",
        "Proxy outcome with incomplete secondary-source, symptom, MBI-LCBI, and infection-prevention adjudication.",
        "procedureevents-derived line exposure is positive documentation evidence, not a complete denominator.",
        "The 2020-2022 period is a post-hoc historical sensitivity cohort because it was repeatedly inspected.",
        "Modest discrimination, near-zero Brier skill, wide calibration intervals, and substantial false-review burden.",
        "Subgroup estimates are descriptive and often sparse; they are not formal fairness or causal results.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("Minimum Requirements Before Clinical Use", level=1)
    for text in [
        "Independent longitudinal EHR validation with complete line episodes, microbiology, daily predictors, and competing-event handling.",
        "Prospectively specified recalibration and review capacity, followed by silent workflow evaluation.",
        "Outcome adjudication or validated source-attribution rules on a representative sample.",
        "Human-factors assessment showing acceptable infection-prevention workload without nurse alarm burden.",
        "Governance for drift, missingness, subgroup monitoring, and version control.",
    ]:
        add_numbered(doc, text)

    doc.add_heading("Key Sources", level=1)
    for source in [
        "CDC/NHSN. Bloodstream Infection Event (Central Line-Associated Bloodstream Infection and Non-central Line Associated Bloodstream Infection), January 2026. https://www.cdc.gov/nhsn/pdfs/pscmanual/4psc_clabscurrent.pdf",
        "Johnson AEW et al. MIMIC-IV documentation: microbiologyevents. https://mimic.mit.edu/docs/IV/modules/hosp/microbiologyevents.html",
        "Wei Z, Kanjilal S. ARMD-MGB v1.0.0. PhysioNet. https://doi.org/10.13026/2r5k-b955",
        "Pollard TJ et al. eICU Collaborative Research Database v2.0. PhysioNet. https://doi.org/10.13026/C2WM1R",
        "Collins GS et al. TRIPOD+AI statement. BMJ. 2024;385:e078378. https://doi.org/10.1136/bmj-2023-078378",
    ]:
        add_bullet(doc, source)

    path = STAGE / "v0_5_run33_model_card.docx"
    doc.save(path)
    return path


def build_manuscript(overall, policy, ablation, windows, external_meta, cohort, figures) -> Path:
    doc = Document()
    configure_document(doc, "CVCML v0.5 Publication-Oriented Results Package")
    add_title(doc, "Seven-Day CVC-Associated BSI Proxy Risk Stratification in MIMIC-IV", "Leakage-audited results and external label-component transportability | Run 33")
    add_callout(doc, "Technical summary", "The leakage-safe model retained modest ranking value after removal of an outcome-derived culture feature, but probabilistic skill was minimal and uncertainty was wide. The defensible use case is bounded infection-prevention review prioritization. ARMD-MGB supports transportability of organism and partial source-screen logic, not external validation of the prediction model.")

    doc.add_heading("Abstract-Style Summary", level=1)
    doc.add_heading("Background", level=2)
    doc.add_paragraph(
        "Central-line bloodstream infection surveillance is clinically important but difficult to reproduce from deidentified EHR data. This project developed a daily seven-day risk model using MIMIC-IV and progressively audited catheter episodes, outcome definitions, temporal leakage, calibration, workflow burden, and external label transportability."
    )
    doc.add_heading("Methods", level=2)
    doc.add_paragraph(
        "Recorded catheter procedures were reconstructed into continuous exposure episodes. Daily landmarks used a strict primary-or-uncertain CVC-associated BSI proxy rather than adjudicated NHSN CLABSI. A frozen XGBoost candidate was trained in 2008-2013, Platt-calibrated in 2014-2016, and evaluated in 2017-2019 after excluding the outcome-derived early_positive_culture feature. Patient-clustered bootstrap intervals used 2,000 replicates. External feasibility was assessed in eICU-CRD, and organism/source-rule transportability was characterized in ARMD-MGB."
    )
    doc.add_heading("Results", level=2)
    doc.add_paragraph(
        "The leakage-safe evaluation contained 5,694 landmarks from 648 episodes and 590 patients, with 242 positive landmarks (4.25%) and 54 positive episodes. ROC-AUC was 0.612 (95% CI 0.518-0.703), PR-AUC 0.065 (0.041-0.113), and PR-AUC lift 1.54x prevalence (1.13-2.48x). Brier Skill Score was 0.005 (-0.019 to 0.021). Reviewing the top 10% of episodes yielded 15.4% PPV and 18.5% positive-episode recall. ARMD-MGB contained 32,887 positive blood-culture accessions; 25,940 met the finalized organism rule. eICU failed the exact-validation feasibility gate."
    )
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "Seven-day risk ranking may support bounded retrospective infection-prevention review, but current evidence does not support interruptive alerts, autonomous diagnosis, or absolute-risk interpretation. The study provides leakage-audited internal evidence and external validation of label components, not external institutional validation of model performance."
    )

    doc.add_heading("Study Objective and Claim Boundary", level=1)
    doc.add_paragraph(
        "The publication-oriented question is whether routinely available MIMIC-IV information can prioritize catheter episodes for review of a seven-day strict CVC-associated BSI proxy under realistic workflow constraints. The target is deliberately described as a proxy because structured data cannot fully reproduce NHSN primary/secondary attribution, common-commensal symptom rules, MBI-LCBI exclusions, and infection-prevention adjudication."
    )
    add_callout(doc, "Primary claim", "A leakage-audited model provides modest retrospective ranking enrichment for a bounded infection-prevention review list in MIMIC-IV.")
    add_callout(doc, "Not claimed", "NHSN-adjudicated CLABSI prediction, externally validated clinical performance, or readiness for bedside deployment.", fill="FCE8E6", accent=RED)

    doc.add_heading("Methods", level=1)
    doc.add_heading("Catheter Exposure Episodes", level=2)
    cohort_map = {row.measure: int(row.value) for row in cohort.itertuples()}
    add_table(doc, ["Cohort step", "Count", "Role"], [
        ["Raw recorded CVC procedure events", f"{cohort_map['Raw CVC procedure events']:,}", "Positive documentation evidence"],
        ["Continuous exposure periods", f"{cohort_map['Continuous exposure periods']:,}", "All reconstructed episodes; no longest-line selection"],
        ["Eligible periods >=48 h", f"{cohort_map['Eligible exposure periods (>=48 h)']:,}", "Episode eligibility"],
        ["Stays with multiple periods", f"{cohort_map['Stays with multiple exposure periods']:,}", "Retained rather than discarded"],
        ["Strict proxy-positive eligible episodes", f"{cohort_map['Strict proxy-positive eligible episodes']:,}", "Pre-landmark outcome inventory"],
    ], [2.5, 1.0, 3.3], font_size=8.5)
    doc.add_paragraph(
        "Episode reconstruction corrected the earlier vulnerability of selecting the eventual longest line, which used future information and could misattribute a culture to the wrong exposure. procedureevents remains an incomplete denominator: recorded events imply line evidence, but missing events do not prove absence."
    )

    doc.add_heading("Outcome and Landmark Framing", level=2)
    add_bullet(doc, "Daily landmarks within eligible catheter episodes.")
    add_bullet(doc, "One clinically interpretable horizon: strict primary-or-uncertain proxy event within the next seven days.")
    add_bullet(doc, "Competing-event censoring at line end, discharge, death, or observation end where encoded by the pipeline.")
    add_bullet(doc, "Recognized pathogens qualify directly; common commensals require repeated-culture logic.")
    add_bullet(doc, "Urine/respiratory matching provides partial secondary-source evidence; uncertain events remain separate from clearly secondary-possible events.")

    doc.add_heading("Model and Evaluation", level=2)
    doc.add_paragraph(
        "The frozen XGBoost specification was inherited from the pre-audit development pipeline. Run 29 refit only the prespecified feature variants needed for the leakage ablation and did not tune on 2017-2019. Platt calibration was learned in 2014-2016. Run 30 characterized the safe variant on 2017-2019 using patient-level cluster bootstrap resampling so all landmarks and episodes for a sampled patient moved together."
    )
    add_table(doc, ["Period", "Role", "Status"], [
        ["2008-2013", "Model development", "Used for frozen candidate"],
        ["2014-2016", "Platt calibration", "Calibration map fixed before validation"],
        ["2017-2019", "Leakage-safe characterization", "Primary Run 30 evidence"],
        ["2020-2022", "Historical temporal sensitivity", "Repeatedly inspected; not a pristine safe-model lockbox"],
    ], [1.2, 2.35, 3.25], font_size=8.5)

    doc.add_heading("Leakage Audit Results", level=1)
    original = ablation[(ablation.model_variant == "original_episode_level_early_positive") & (ablation.split == "validation") & (ablation.calibration == "platt")].iloc[0]
    safe = ablation[(ablation.model_variant == "safe_exclude_early_positive") & (ablation.split == "validation") & (ablation.calibration == "platt")].iloc[0]
    doc.add_paragraph(
        "The early_positive_culture variable was copied across episode landmarks and encoded eventual culture positivity before results were necessarily available. In development, 42.4% of flagged landmark rows lacked an organism-positive storetime by the landmark. Removing the feature changed validation PR-AUC from "
        f"{original.pr_auc:.4f} to {safe.pr_auc:.4f} (difference {safe.pr_auc-original.pr_auc:.4f}) and therefore removed a validity threat without materially changing ranking performance."
    )
    add_figure(doc, figures["result_availability"], "Figure 1. Result-availability audit showing why early_positive_culture was excluded from the manuscript-safe model.", width=6.2)

    doc.add_heading("Leakage-Safe Performance", level=1)
    roc = metric_row(overall, "roc_auc")
    pr = metric_row(overall, "pr_auc")
    prevalence = metric_row(overall, "prevalence")
    lift = metric_row(overall, "pr_auc_lift")
    brier = metric_row(overall, "brier_score")
    bss = metric_row(overall, "brier_skill_score")
    intercept = metric_row(overall, "calibration_intercept")
    slope = metric_row(overall, "calibration_slope")
    eo = metric_row(overall, "expected_observed_ratio")
    add_table(doc, ["Metric", "Estimate", "95% CI", "Reading"], [
        ["Row prevalence", f"{prevalence.estimate:.2%}", f"{prevalence.ci_lower_95:.2%}-{prevalence.ci_upper_95:.2%}", "Baseline for PR-AUC"],
        ["ROC-AUC", f"{roc.estimate:.3f}", f"{roc.ci_lower_95:.3f}-{roc.ci_upper_95:.3f}", "Modest discrimination"],
        ["PR-AUC", f"{pr.estimate:.3f}", f"{pr.ci_lower_95:.3f}-{pr.ci_upper_95:.3f}", "Low absolute precision-recall area"],
        ["PR-AUC lift", f"{lift.estimate:.2f}x", f"{lift.ci_lower_95:.2f}-{lift.ci_upper_95:.2f}x", "Above prevalence baseline"],
        ["Brier score", f"{brier.estimate:.3f}", f"{brier.ci_lower_95:.3f}-{brier.ci_upper_95:.3f}", "Must be compared with prevalence model"],
        ["Brier Skill Score", f"{bss.estimate:.3f}", f"{bss.ci_lower_95:.3f}-{bss.ci_upper_95:.3f}", "No clear probabilistic improvement"],
        ["Calibration intercept", f"{intercept.estimate:.3f}", f"{intercept.ci_lower_95:.3f}-{intercept.ci_upper_95:.3f}", "Average calibration uncertain"],
        ["Calibration slope", f"{slope.estimate:.3f}", f"{slope.ci_lower_95:.3f}-{slope.ci_upper_95:.3f}", "Wide CI; overfitting cannot be excluded"],
        ["Expected:observed", f"{eo.estimate:.3f}", f"{eo.ci_lower_95:.3f}-{eo.ci_upper_95:.3f}", "Tendency toward overprediction"],
    ], [1.55, 1.05, 1.7, 2.5], font_size=8.1)
    doc.add_paragraph(
        "The safe model is better understood as a weak-to-modest ranker than as an absolute risk model. Its PR-AUC exceeds prevalence, but the Brier Skill Score is approximately zero and its confidence interval includes negative values."
    )
    add_figure(doc, figures["calibration"], "Figure 2. Calibration by risk group in the 2017-2019 leakage-safe evaluation.", width=6.2)

    doc.add_heading("Operational Characterization", level=1)
    rows = []
    for top in [1, 5, 10, 20]:
        n = policy_value(policy, top, "episodes_reviewed")
        ppv = policy_value(policy, top, "episode_ppv")
        recall = policy_value(policy, top, "episode_recall")
        false = policy_value(policy, top, "false_reviews_per_true_positive")
        rows.append([
            f"Top {top}%", int(round(n.estimate)),
            f"{ppv.estimate:.1%} ({ppv.ci_lower_95:.1%}-{ppv.ci_upper_95:.1%})",
            f"{recall.estimate:.1%} ({recall.ci_lower_95:.1%}-{recall.ci_upper_95:.1%})",
            f"{false.estimate:.2f} ({false.ci_lower_95:.2f}-{false.ci_upper_95:.2f})",
        ])
    add_table(doc, ["Budget", "Episodes", "PPV (95% CI)", "Recall (95% CI)", "False reviews/TP (95% CI)"], rows,
              [1.0, 0.8, 1.75, 1.65, 1.6], font_size=7.8)
    doc.add_paragraph(
        "A fixed-capacity review list is the most plausible operational framing. The top 10% policy reviewed 65 of 648 episodes, captured 10 of 54 positive episodes, and produced 55 false-positive reviews. This workload could be acceptable for periodic infection-prevention review but is inappropriate for high-frequency nurse-facing alerts."
    )
    add_figure(doc, figures["review_policy"], "Figure 3. Episode-level review yield and false-review burden across fixed review budgets.", width=6.2)

    doc.add_heading("Subgroup Characterization", level=1)
    doc.add_paragraph(
        "Sex, age, race, ICU type, admission type, insurance, and catheter context were prespecified descriptive subgroup checks. Seven levels met a stability rule of at least 100 patients and 20 positive patients. Sparse cells had confidence intervals suppressed. No multiplicity-adjusted tests were performed, so results should not be interpreted as formal fairness claims or causal differences."
    )
    add_figure(doc, figures["subgroups"], "Figure 4. Descriptive PR-AUC lift across subgroups with sufficient event support.", width=6.2)

    doc.add_heading("External Validation Feasibility and Label Transportability", level=1)
    doc.add_heading("eICU-CRD", level=2)
    doc.add_paragraph(
        "eICU includes multicenter ICU predictors, but sparse microbiology and incomplete explicit line timing prevented reconstruction of the frozen task. Only one positive-blood stay met the prespecified criterion of at least 48 hours after explicit line placement. Reporting an external AUROC or PR-AUC from this selected subset would be misleading, so no model-performance estimate was produced."
    )
    doc.add_heading("ARMD-MGB", level=2)
    total = int(external_meta["armd_positive_blood_accessions"])
    strict = int(external_meta["armd_strict_accessions"])
    add_table(doc, ["External component", "Result", "Interpretation"], [
        ["Positive blood-culture accessions", f"{total:,}", "Culture-positive denominator, not CLABSI prevalence"],
        ["Final strict organism rule", f"{strict:,} ({strict/total:.1%})", "Rule is implementable at another health system"],
        ["Recognized pathogen", f"{int(external_meta['armd_recognized_pathogen_accessions']):,}", "Direct organism qualification"],
        ["Repeated common commensal", f"{int(external_meta['armd_repeated_commensal_48h_accessions']):,}", "Separate same/adjacent-day accessions"],
        ["Shared organism categories", f"{int(external_meta['shared_organisms'])}", f"Rank rho={float(external_meta['shared_organism_spearman_rho']):.3f}; JSD={float(external_meta['jensen_shannon_divergence']):.3f}"],
    ], [2.15, 1.45, 3.2], font_size=8.4)
    doc.add_paragraph(
        "ARMD-MGB validates the portability of organism classification and quantifies partial urine/respiratory source evidence. It lacks catheter exposure episodes and daily predictors, so it cannot validate discrimination, calibration, or review burden."
    )
    add_figure(doc, figures["organism_profile"], "Figure 5. Canonical organism distributions in MIMIC-IV and ARMD-MGB, demonstrating moderate rank agreement and meaningful case-mix shift.", width=6.2)

    source_rows = []
    for row in windows.itertuples():
        source_rows.append([
            f"+/-{int(row.window_days)} day",
            f"{int(row.matching_source_accessions):,}/{int(row.strict_accessions):,}",
            f"{row.proportion:.1%}",
            f"{row.ci_low:.1%}-{row.ci_high:.1%}",
        ])
    add_table(doc, ["Window", "Same-organism source evidence", "Proportion", "Wilson 95% CI"], source_rows,
              [1.1, 2.4, 1.25, 1.55], font_size=8.4)
    doc.add_paragraph(
        "Same-organism urine or respiratory cultures are evidence of a possible secondary source, not proof of NHSN secondary-BSI attribution. Wound, abdominal, symptom, MBI-LCBI, and infection-prevention adjudication remain unavailable."
    )
    add_figure(doc, figures["source_window"], "Figure 6. Sensitivity of partial secondary-source evidence to the prespecified culture-matching window.", width=6.2)

    doc.add_heading("Interpretation", level=1)
    add_bullet(doc, "The strongest contribution is methodological honesty: catheter episodes, temporal landmarks, outcome terminology, leakage, calibration, workflow burden, and external feasibility were explicitly audited.")
    add_bullet(doc, "Performance is below the aspirational publication target of AUROC >=0.70 and does not provide a twofold PR-AUC lift with stable calibration.")
    add_bullet(doc, "The score may still be useful as a bounded ranker because top-decile episode PPV is about 1.85 times the episode prevalence (15.4% versus 8.3%).")
    add_bullet(doc, "Near-zero Brier skill and wide calibration uncertainty rule out bedside absolute-risk claims.")
    add_bullet(doc, "External component-level outcome evidence is valuable but cannot substitute for a longitudinal external cohort with line episodes and predictors.")

    doc.add_heading("Limitations", level=1)
    for text in [
        "The target is a strict CVC-associated BSI proxy, not adjudicated NHSN CLABSI.",
        "Line ascertainment is incomplete when procedure documentation is absent.",
        "The current primary evaluation is a development-validation period, and the later MIMIC interval is no longer a pristine lockbox.",
        "Confidence intervals reflect sampling uncertainty conditional on the fitted model, not the full development process.",
        "External databases did not permit full prediction-model validation.",
        "ARMD-MGB comparisons use culture accessions from broader care settings, not eligible ICU catheter episodes.",
        "Subgroup analyses are descriptive and underpowered for some strata.",
        "Prospective workflow effects, clinician response, and patient outcomes were not evaluated.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("Publication-Safe Conclusions", level=1)
    add_callout(doc, "Supported", "In MIMIC-IV, a leakage-audited seven-day model offered modest enrichment for bounded infection-prevention review of a strict CVC-associated BSI proxy. External microbiology data supported transportability of organism and partial source-screen logic.", fill="E7F3EC", accent=GREEN)
    add_callout(doc, "Unsupported", "Claims of NHSN CLABSI prediction, clinically calibrated absolute risk, external model validation, nurse-alert readiness, or improved patient outcomes.", fill="FCE8E6", accent=RED)

    doc.add_heading("Recommended Next Study", level=1)
    for text in [
        "Freeze Run 33 as the reporting baseline and stop broad tuning on repeatedly inspected MIMIC periods.",
        "Obtain a longitudinal external hospital EHR with central-line start/end, microbiology, daily predictors, discharge/death/removal, and adequate positive episodes.",
        "Prespecify variable mapping, missingness handling, outcome reconstruction, recalibration policy, and review budgets before outcome-linked scoring.",
        "Perform external discrimination, calibration, Brier Skill Score, and top-k review-burden analysis with patient/episode clustered confidence intervals.",
        "If full external data remain unavailable, complete blinded note/source adjudication and present the work as a transparent single-center methods and feasibility study.",
    ]:
        add_numbered(doc, text)

    doc.add_heading("References", level=1)
    references = [
        "Centers for Disease Control and Prevention. Bloodstream Infection Event (Central Line-Associated Bloodstream Infection and Non-central Line Associated Bloodstream Infection). NHSN Patient Safety Component Manual, January 2026. https://www.cdc.gov/nhsn/pdfs/pscmanual/4psc_clabscurrent.pdf",
        "Johnson AEW, Bulgarelli L, Shen L, et al. MIMIC-IV, a freely accessible electronic health record dataset. Scientific Data. 2023;10:1. https://doi.org/10.1038/s41597-022-01899-x",
        "MIMIC-IV documentation. microbiologyevents table. https://mimic.mit.edu/docs/IV/modules/hosp/microbiologyevents.html",
        "Wei Z, Kanjilal S. Antibiotic Resistance Microbiology Dataset Mass General Brigham (ARMD-MGB), version 1.0.0. PhysioNet. 2025. https://doi.org/10.13026/2r5k-b955",
        "Pollard TJ, Johnson AEW, Raffa JD, Celi LA, Badawi O, Mark RG. eICU Collaborative Research Database, version 2.0. PhysioNet. 2019. https://doi.org/10.13026/C2WM1R",
        "Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378. https://doi.org/10.1136/bmj-2023-078378",
    ]
    for ref in references:
        add_numbered(doc, ref)

    path = STAGE / "v0_5_run33_manuscript_results_package.docx"
    doc.save(path)
    return path


def write_markdown(overall, policy, external_meta) -> None:
    roc = metric_row(overall, "roc_auc")
    pr = metric_row(overall, "pr_auc")
    bss = metric_row(overall, "brier_skill_score")
    p10_ppv = policy_value(policy, 10, "episode_ppv")
    p10_rec = policy_value(policy, 10, "episode_recall")
    model_card = f"""# CVCML v0.5 Leakage-Safe Model Card

**Status:** Research prototype; not for clinical use.

## Intended use

Bounded infection-prevention review-list prioritization for seven-day strict primary-or-uncertain CVC-associated BSI proxy risk. Not an autonomous diagnosis, calibrated bedside risk tool, or interruptive alert.

## Frozen safe candidate

- XGBoost with Platt calibration.
- Train 2008-2013; calibration 2014-2016; leakage-safe evaluation 2017-2019.
- `early_positive_culture` excluded after Run 29 outcome-validity audit.
- Target: `future_strict_primary_or_uncertain_cvc_bsi_proxy_7d`.

## Performance

- ROC-AUC: {roc.estimate:.3f} (95% CI {roc.ci_lower_95:.3f}-{roc.ci_upper_95:.3f}).
- PR-AUC: {pr.estimate:.3f} ({pr.ci_lower_95:.3f}-{pr.ci_upper_95:.3f}); prevalence 4.25%.
- Brier Skill Score: {bss.estimate:.3f} ({bss.ci_lower_95:.3f}-{bss.ci_upper_95:.3f}).
- Top 10% episode review PPV: {p10_ppv.estimate:.1%}; recall: {p10_rec.estimate:.1%}.

## External evidence

- ARMD-MGB: {int(external_meta['armd_positive_blood_accessions']):,} positive blood-culture accessions; organism and partial source logic transported.
- eICU: failed exact-validation feasibility gate.
- Full external model validation: not achieved.

## Decision

Credible as a transparent retrospective review-prioritization research model. Not ready for clinical deployment.
"""
    (STAGE / "v0_5_run33_model_card.md").write_text(model_card, encoding="utf-8")

    manuscript = f"""# Run 33 Manuscript Results Summary

The leakage-safe candidate was characterized on 5,694 daily landmarks from 648 catheter episodes and 590 patients in 2017-2019. There were 242 positive landmarks (4.25%) and 54 positive episodes. ROC-AUC was {roc.estimate:.3f} (95% CI {roc.ci_lower_95:.3f}-{roc.ci_upper_95:.3f}), PR-AUC was {pr.estimate:.3f} ({pr.ci_lower_95:.3f}-{pr.ci_upper_95:.3f}), and Brier Skill Score was {bss.estimate:.3f} ({bss.ci_lower_95:.3f}-{bss.ci_upper_95:.3f}).

Run 29 removed `early_positive_culture` because it encoded eventual culture positivity before results were necessarily available. The feature exclusion changed validation PR-AUC by only -0.0038. At a top-10% review budget, PPV was {p10_ppv.estimate:.1%} and positive-episode recall was {p10_rec.estimate:.1%}.

ARMD-MGB supported external transportability of organism and partial secondary-source logic, but neither ARMD-MGB nor eICU could reproduce the full catheter-episode prediction task. The project therefore has leakage-audited internal evidence and external label-component validation, not external model validation.

The manuscript-safe conclusion is that the model offers modest retrospective enrichment for a bounded infection-prevention review list targeting a strict CVC-associated BSI proxy. It does not support NHSN CLABSI claims, interruptive alerts, or absolute-risk use.
"""
    (STAGE / "v0_5_run33_manuscript_results.md").write_text(manuscript, encoding="utf-8")


def append_hub(overall, policy, external_meta, figures) -> Path:
    doc = Document(HUB_SOURCE)
    doc.add_page_break()
    doc.add_heading("Run 33 - Publication Consolidation and Final Project Status", level=1)
    doc.add_paragraph(
        "Purpose: supersede the provisional Run 28 reporting package with leakage-safe results, patient-clustered uncertainty, operational review-burden evidence, and the external feasibility and label-transportability findings from Runs 31-32. Run 33 is a reporting and evidence-synthesis run; it does not refit or retune the model."
    )

    doc.add_heading("What Run 33 Changed", level=2)
    add_bullet(doc, "Replaced the provisional Run 28 model card and manuscript summary with the Run 29 leakage-safe candidate.")
    add_bullet(doc, "Added exact Run 30 confidence intervals, Brier Skill Score, calibration, subgroup, and review-budget interpretation.")
    add_bullet(doc, "Reframed the 2020-2022 MIMIC period as a post-hoc historical sensitivity because it was repeatedly inspected before the safe model was finalized.")
    add_bullet(doc, "Added Run 31 eICU failed-feasibility evidence and Run 32 ARMD-MGB external label-component transportability.")
    add_bullet(doc, "Created a claim register that separates supported, cautiously supported, and unsupported statements.")

    roc = metric_row(overall, "roc_auc")
    pr = metric_row(overall, "pr_auc")
    lift = metric_row(overall, "pr_auc_lift")
    bss = metric_row(overall, "brier_skill_score")
    slope = metric_row(overall, "calibration_slope")
    p10_ppv = policy_value(policy, 10, "episode_ppv")
    p10_rec = policy_value(policy, 10, "episode_recall")
    add_table(doc, ["Evidence domain", "Run 33 result", "Interpretation"], [
        ["Evaluation cohort", "5,694 landmarks; 648 episodes; 590 patients; 54 positive episodes", "2017-2019 leakage-safe characterization"],
        ["Discrimination", f"ROC-AUC {roc.estimate:.3f} ({roc.ci_lower_95:.3f}-{roc.ci_upper_95:.3f}); PR-AUC {pr.estimate:.3f} ({pr.ci_lower_95:.3f}-{pr.ci_upper_95:.3f})", "Modest and imprecise"],
        ["Lift", f"{lift.estimate:.2f}x prevalence ({lift.ci_lower_95:.2f}-{lift.ci_upper_95:.2f}x)", "Some ranking enrichment"],
        ["Probability quality", f"Brier Skill {bss.estimate:.3f}; calibration slope {slope.estimate:.3f}", "Absolute-risk use not supported"],
        ["Top 10% review", f"PPV {p10_ppv.estimate:.1%}; recall {p10_rec.estimate:.1%}", "Potential bounded review queue, not alert"],
        ["External outcome components", f"ARMD-MGB: {int(external_meta['armd_positive_blood_accessions']):,} positive blood accessions", "Organism/source logic transported"],
        ["External full-model validation", "Not achieved", "eICU failed feasibility; ARMD-MGB lacks line/predictor timelines"],
    ], [1.65, 2.35, 2.8], font_size=8.2)

    doc.add_heading("Final Model Interpretation", level=2)
    doc.add_paragraph(
        "The current candidate is a modest retrospective ranker for infection-prevention review, not a bedside risk calculator. Its strongest operational result is enrichment within a fixed-capacity episode review list. Near-zero Brier skill, wide calibration intervals, and several false reviews per true positive prevent an absolute-risk or interruptive-alert claim."
    )
    add_figure(doc, figures["review_policy"], "Run 33. Review yield and workload across fixed episode-review budgets.", width=6.1)

    doc.add_heading("External Evidence Boundary", level=2)
    doc.add_paragraph(
        "ARMD-MGB demonstrated that the organism rule and partial urine/respiratory source screen can be applied in another health system, with moderate organism-rank agreement and meaningful distribution shift. eICU did not contain enough microbiology-linked explicit catheter episodes for honest model validation. No external AUROC, PR-AUC, calibration, or review-burden estimate is claimed."
    )
    add_figure(doc, figures["organism_profile"], "Run 33. MIMIC-IV and ARMD-MGB organism profiles after canonical harmonization.", width=6.1)

    doc.add_heading("Final Roadmap Status", level=2)
    add_table(doc, ["Roadmap component", "Status", "Evidence / remaining work"], [
        ["Cohort and catheter episodes", "Complete for MIMIC-IV", "All recorded exposure episodes retained; longest-line selection removed"],
        ["Seven-day dynamic landmarks", "Complete", "Daily landmark task with competing-event censoring logic"],
        ["Proxy-label hierarchy", "Complete within structured data", "Broad, strict primary, secondary-possible, and uncertain outcomes"],
        ["Leakage audit", "Complete", "Outcome-derived culture feature excluded"],
        ["Internal characterization", "Complete", "Clustered CIs, calibration, Brier skill, subgroups, review burden"],
        ["Temporal evidence", "Available with caveat", "Later MIMIC period is historical sensitivity, not pristine lockbox"],
        ["External label-component evidence", "Complete", "ARMD-MGB organism and partial source-rule transportability"],
        ["External model validation", "Outstanding", "Requires longitudinal EHR with line and predictor timelines"],
        ["Publication package", "Complete through Run 33", "Superseding model card, manuscript package, claim register, evidence traceability"],
        ["Clinical deployment", "Not ready", "Requires external validation, prospective workflow study, and governance"],
    ], [2.0, 1.35, 3.45], font_size=8.0)

    doc.add_heading("Publication Position After Run 33", level=2)
    add_callout(doc, "Defensible manuscript contribution", "A transparent MIMIC-IV replication and methods study showing that seven-day strict CVC-associated BSI proxy risk can be modestly enriched for bounded review, while leakage, calibration, alert burden, label uncertainty, and external-data feasibility materially constrain clinical claims.", fill="E7F3EC", accent=GREEN)
    add_callout(doc, "Main unresolved requirement", "Full external institutional validation of the frozen model on a longitudinal EHR with complete catheter episodes, microbiology, daily predictors, and competing-event information.", fill="FFF4CE", accent=GOLD)

    doc.add_heading("Run 33 Deliverables", level=2)
    for name in [
        "v0_5_run33_model_card.docx / .md",
        "v0_5_run33_manuscript_results_package.docx / .md",
        "v0_5_run33_performance_summary.csv",
        "v0_5_run33_review_policy_summary.csv",
        "v0_5_run33_external_validation_summary.csv",
        "v0_5_run33_claim_register.csv",
        "v0_5_run33_evidence_traceability.csv",
        "v0_5_run33_limitations_and_actions.csv",
        "v0_5_run33_manifest.json",
    ]:
        add_bullet(doc, name)

    doc.add_heading("Run 33 Sources", level=2)
    for source in [
        "CDC/NHSN Bloodstream Infection Event chapter, January 2026: https://www.cdc.gov/nhsn/pdfs/pscmanual/4psc_clabscurrent.pdf",
        "MIMIC-IV microbiologyevents documentation: https://mimic.mit.edu/docs/IV/modules/hosp/microbiologyevents.html",
        "ARMD-MGB v1.0.0: https://doi.org/10.13026/2r5k-b955",
        "eICU-CRD v2.0: https://doi.org/10.13026/C2WM1R",
        "TRIPOD+AI reporting guidance: https://doi.org/10.1136/bmj-2023-078378",
    ]:
        add_bullet(doc, source)

    doc.add_heading("Next Step", level=2)
    doc.add_paragraph(
        "Freeze Run 33 as the current manuscript/reporting baseline. The next scientific run should begin only when a suitable longitudinal external EHR or a blinded source-adjudication resource is available. Broad retuning on the repeatedly inspected MIMIC periods is no longer the highest-value activity."
    )

    doc.save(HUB_OUTPUT)
    return HUB_OUTPUT


def validate_docx(path: Path, required: list[str]) -> dict:
    doc = Document(path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            all_text += "\n" + " | ".join(cell.text for cell in row.cells)
    missing = [term for term in required if term not in all_text]
    if missing:
        raise RuntimeError(f"{path.name} missing required text: {missing}")
    return {
        "file": path.name,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "missing_required_text": missing,
    }


def write_manifest(files: list[Path], qa: list[dict]) -> None:
    manifest = {
        "run": 33,
        "name": "v0.5 Publication Consolidation",
        "created": date.today().isoformat(),
        "model_refit": False,
        "model_retuned": False,
        "primary_model_variant": "safe_exclude_early_positive",
        "primary_evaluation_period": "2017-2019 development-validation",
        "full_external_model_validation": False,
        "external_label_component_validation": True,
        "claims_status": "publication-oriented retrospective methods/results package; not clinical deployment",
        "qa": qa,
        "files": [],
    }
    for path in files:
        if path.exists():
            manifest["files"].append({
                "name": path.relative_to(STAGE).as_posix() if STAGE in path.parents else path.name,
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
            })
    with (STAGE / "v0_5_run33_manifest.json").open("w", encoding="utf-8") as handle:
        json.dump(manifest, handle, indent=2)


def main() -> None:
    STAGE.mkdir(parents=True, exist_ok=True)
    figures = copy_figures()
    overall, policy, ablation, windows, external_scope, external_meta, cohort = load_data()
    write_csvs(overall, policy, ablation, windows, external_meta, cohort)
    write_markdown(overall, policy, external_meta)
    model_card = build_model_card(overall, policy, external_meta, figures)
    manuscript = build_manuscript(overall, policy, ablation, windows, external_meta, cohort, figures)
    hub = append_hub(overall, policy, external_meta, figures)

    notes = """# Run 33 - Publication Consolidation

Run 33 is a reporting and evidence-synthesis run. It does not refit, retune, or rescore the 2020-2022 MIMIC period.

## Final interpretation

- Manuscript-safe model: Run 29 `safe_exclude_early_positive`, Platt calibrated.
- Primary evidence: Run 30 patient-clustered characterization in 2017-2019.
- Intended use: bounded infection-prevention review prioritization.
- Outcome: strict primary-or-uncertain CVC-associated BSI proxy, not adjudicated NHSN CLABSI.
- External evidence: ARMD-MGB organism/source component transportability; eICU failed full-model feasibility.
- External full-model validation: outstanding.

## Decision

Freeze Run 33 as the current publication baseline. Do not perform broad additional tuning on repeatedly inspected MIMIC periods. Prioritize longitudinal external validation or blinded source/notes adjudication.
"""
    (STAGE / "v0_5_run33_notes.md").write_text(notes, encoding="utf-8")

    qa = [
        validate_docx(model_card, ["Intended Use", "Leakage Audit", "External Evidence", "Minimum Requirements Before Clinical Use"]),
        validate_docx(manuscript, ["Abstract-Style Summary", "Leakage-Safe Performance", "External Validation Feasibility", "Publication-Safe Conclusions"]),
        validate_docx(hub, ["Run 33 - Publication Consolidation and Final Project Status", "Final Roadmap Status", "Publication Position After Run 33"]),
    ]

    files = [p for p in STAGE.rglob("*") if p.is_file()]
    write_manifest(files, qa)
    print(json.dumps({
        "stage": str(STAGE),
        "model_card": str(model_card),
        "manuscript": str(manuscript),
        "hub": str(hub),
        "qa": qa,
    }, indent=2))


if __name__ == "__main__":
    main()

