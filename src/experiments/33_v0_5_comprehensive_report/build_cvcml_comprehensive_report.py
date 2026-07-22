from __future__ import annotations

from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\path\to\CVCML")
OUT = Path(r"C:\path\to\CVCML\Reports")
OUT.mkdir(parents=True, exist_ok=True)

REPORT_DOCX = OUT / "CVCML Comprehensive Project Report - Runs 1-33.docx"

ROADMAP = Path(r"C:\path\to\clabsi_project_roadmap.png")
RUN33 = ROOT / "Outputs" / "Run 33 (v0.5 Publication Consolidation)"
FIGURES = {
    "leakage": RUN33 / "figures" / "v0_5_run29_result_unavailability_by_landmark.png",
    "calibration": RUN33 / "figures" / "v0_5_run30_calibration_deciles.png",
    "policy": RUN33 / "figures" / "v0_5_run30_episode_review_policy.png",
    "subgroups": RUN33 / "figures" / "v0_5_run30_subgroup_pr_auc_lift.png",
    "organisms": RUN33 / "figures" / "run32_mimic_armd_organism_profile.png",
    "source_window": RUN33 / "figures" / "run32_secondary_source_window_sensitivity.png",
}


BLUE = "2E75B6"
DARK_BLUE = "17365D"
INK = "1F2937"
MUTED = "5B6573"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF4EA"
PALE_GOLD = "FFF4CE"
PALE_RED = "FDECEC"
LIGHT_GRAY = "F2F4F7"
BORDER = "CBD5E1"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_widths(table, widths_in):
    table.autofit = False
    total_dxa = int(sum(widths_in) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width)
            set_cell_margins(cell)


def set_run(run, size=10.5, bold=False, color=INK, italic=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, code):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, fld_end])


def add_para(doc, text="", size=10.5, bold=False, color=INK, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.10,
             keep=False, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_rich_para(doc, parts, before=0, after=6, line=1.10, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, options in parts:
        set_run(p.add_run(text), **options)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text),
            size={1: 16, 2: 13, 3: 11.5}[level],
            bold=True,
            color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=10.5)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, label_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    set_run(p.add_run(f"{label}: "), size=10.5, bold=True, color=label_color)
    set_run(p.add_run(text), size=10.5, color=INK)
    add_para(doc, "", after=3)
    return table


def add_table(doc, headers, rows, widths, font_size=9.2, header_fill=BLUE,
              first_col_bold=False, zebra=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_cant_split(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_run(p.add_run(str(header)), size=font_size, bold=True,
                color=WHITE if header_fill == BLUE else DARK_BLUE)
    for r_idx, row in enumerate(rows):
        new_row = table.add_row()
        set_cant_split(new_row)
        cells = new_row.cells
        if zebra and r_idx % 2:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for c_idx, value in enumerate(row):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run(p.add_run(str(value)), size=font_size,
                    bold=first_col_bold and c_idx == 0,
                    color=DARK_BLUE if first_col_bold and c_idx == 0 else INK)
    return table


def add_caption(doc, text):
    p = add_para(doc, text, size=9, color=MUTED, italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=9, line=1.0)
    p.paragraph_format.keep_with_next = False
    return p


def add_figure(doc, path: Path, caption: str, width=6.2):
    if not path.exists():
        add_callout(doc, "Figure unavailable", str(path), fill=PALE_RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def page_break(doc):
    doc.add_page_break()


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, before, after, color in [
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 11.5, 8, 4, DARK_BLUE),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Bullet 2"]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run(hp.add_run("CVCML  |  Comprehensive Project Report"), size=8.5,
            bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(fp.add_run("Research prototype  |  "), size=8.5, color=MUTED)
    add_field(fp, "PAGE")


def add_cover(doc):
    add_para(doc, "CLINICAL MACHINE LEARNING PROJECT", size=10, bold=True,
             color=BLUE, before=34, after=20)
    add_para(doc, "CVCML", size=30, bold=True, color=DARK_BLUE, after=3,
             line=1.0)
    add_para(doc, "Leakage-Audited Prediction of CVC-Associated Bloodstream Infection Risk",
             size=19, bold=True, color=INK, after=9, line=1.10)
    add_para(doc, "A comprehensive account of the clinical premise, data pipeline, model evolution, validation evidence, operational interpretation, and current limitations across Runs 1-33.",
             size=12, color=MUTED, after=30, line=1.25)

    add_table(doc,
              ["Project scope", "Current status", "Primary dataset"],
              [["Static and dynamic risk modeling", "Research prototype", "MIMIC-IV"],
               ["Runs 1-33", "Not for clinical use", "External component checks: ARMD-MGB, eICU"]],
              [2.0, 2.0, 2.3], font_size=9.6, header_fill=DARK_BLUE)
    add_para(doc, "", after=18)
    add_callout(doc, "Current model role",
                "A retrospective ranking tool for bounded infection-prevention review, not an autonomous diagnosis, interruptive bedside alert, or calibrated absolute-risk calculator.",
                fill=PALE_GOLD, label_color=DARK_BLUE)
    add_para(doc, "Prepared for project explanation and presentation", size=10.5,
             color=MUTED, before=28, after=2)
    add_para(doc, "Updated July 17, 2026", size=10.5, bold=True, color=DARK_BLUE)
    page_break(doc)


def add_exec_summary(doc):
    add_heading(doc, "Executive Summary", 1)
    add_para(doc,
        "CVCML is an end-to-end clinical machine-learning project built to investigate whether routinely collected electronic health record data can identify catheter episodes at elevated risk of a central-venous-catheter-associated bloodstream infection within the next seven days. The work began as a static XGBoost reconstruction and evolved into a leakage-audited, episode-based, daily landmark modeling system with explicit outcome uncertainty, competing events, calibration analysis, review-burden characterization, and external label-component testing.")
    add_callout(doc, "The central result",
                "The final leakage-safe model contains real ranking signal, but the signal is modest. Its most defensible use is prioritizing a small infection-prevention review queue. The evidence does not support bedside probability display, interruptive alarms, or a claim of adjudicated NHSN CLABSI prediction.",
                fill=PALE_BLUE)

    add_heading(doc, "Project at a Glance", 2)
    add_table(doc,
              ["Measure", "Current result", "Meaning"],
              [
                  ["Eligible CVC exposure periods", "11,602", "All reconstructed periods lasting at least 48 hours"],
                  ["Final evaluation frame", "5,694 daily landmarks", "648 episodes, 590 patients, 54 positive episodes"],
                  ["ROC-AUC", "0.612 (95% CI 0.518-0.703)", "Modest discrimination with substantial uncertainty"],
                  ["PR-AUC", "0.065 (0.041-0.113)", "1.54x the 4.25% event prevalence"],
                  ["Brier Skill Score", "0.005 (-0.019 to 0.021)", "No reliable improvement in absolute-risk accuracy over prevalence alone"],
                  ["Top 10% episode review", "PPV 15.4%; recall 18.5%", "Useful enrichment, but most reviewed episodes remain false positives"],
                  ["External evidence", "Component-level only", "ARMD-MGB supported organism/source logic; no external full-model validation"],
                  ["Project endpoint", "Run 33 complete", "Leakage-safe evidence package, model card, and bounded claim register"],
              ],
              [1.65, 1.75, 2.9], font_size=9.2, first_col_bold=True)

    add_heading(doc, "What Makes the Project Strong", 2)
    for text in [
        "The pipeline was repeatedly audited for temporal leakage and outcome-dependent reference times rather than preserving attractive but misleading scores.",
        "The cohort was redesigned from a longest-line-per-stay shortcut to all recorded catheter exposure episodes and daily prediction landmarks.",
        "The outcome was renamed and refined as a strict CVC-associated BSI proxy, with organism rules and a partial secondary-source screen.",
        "Evaluation emphasizes PR-AUC lift, calibration, patient-clustered uncertainty, and review workloadâ€”not AUROC alone.",
        "The final clinical framing is deliberately narrow: bounded review prioritization rather than an alarm system.",
    ]:
        add_bullet(doc, text)

    page_break(doc)
    add_heading(doc, "In This Report", 1)
    add_table(doc,
              ["Section", "Question answered"],
              [
                  ["1. Clinical premise", "What problem is the project trying to solve?"],
                  ["2. Data and cohort", "How were catheter episodes, labels, and landmarks constructed?"],
                  ["3. Model evolution", "How did the work move from static modeling to a leakage-safe dynamic model?"],
                  ["4. Final evidence", "What did the frozen candidate actually achieve?"],
                  ["5. Interpretation", "How could the score be used, and where would it be unsafe?"],
                  ["6. External evidence", "What has been tested outside the original modeling frame?"],
                  ["7. Current state", "What is complete, uncertain, and appropriate as future work?"],
                  ["Appendix", "What happened in each of the 33 runs?"],
              ], [2.0, 4.3], font_size=9.2, first_col_bold=True, header_fill=DARK_BLUE)
    add_heading(doc, "How to Read the Evidence", 2)
    add_table(doc,
              ["Rule", "How it is applied"],
              [
                  ["Later runs supersede earlier claims", "The final leakage-safe evidence takes precedence over attractive preliminary scores."],
                  ["Different frames are not a leaderboard", "Static, dynamic, pragmatic-label, strict-label, and source-screened results answer different questions."],
                  ["Ranking is distinct from probability accuracy", "Discrimination and review enrichment can exist even when calibration and Brier skill are weak."],
                  ["Missing evidence is not a negative", "Incomplete line or note documentation is preserved as uncertainty rather than forced into a clean label."],
              ], [2.15, 4.15], font_size=9.2, first_col_bold=True)
    add_callout(doc, "Recommended reading path",
                "For a short overview, read the Executive Summary, Model Evolution, Final Evidence, and Current Project State. The appendix preserves the complete run history.",
                fill=PALE_BLUE)
    page_break(doc)


def add_clinical_premise(doc):
    add_heading(doc, "1. Clinical Premise", 1)
    add_heading(doc, "Why This Problem Matters", 2)
    add_para(doc,
        "Central venous catheters are essential for intensive care but create a pathway for bloodstream infection. Earlier recognition could support line review, culture interpretation, source investigation, and prevention activity. The challenge is that the event is rare, surveillance definitions are complex, measurements are irregular, and the operational cost of false alarms is high.")
    add_para(doc,
        "The project therefore moved away from a simple questionâ€”whether a classifier can separate positive and negative staysâ€”and toward a harder one: whether a temporally honest score can enrich a finite review queue without creating an unacceptable burden for clinicians or infection-prevention staff.")

    add_heading(doc, "Research Question", 2)
    add_callout(doc, "Primary question",
                "Among reconstructed CVC exposure episodes that remain under observation, can data available at a daily landmark identify episodes likely to meet a strict primary-or-uncertain CVC-associated bloodstream-infection proxy within the next seven days?",
                fill=PALE_BLUE)

    add_heading(doc, "Terminology Boundary", 2)
    add_para(doc,
        "The outcome is not called adjudicated CLABSI. NHSN determination requires an eligible line, laboratory-confirmed bloodstream infection criteria, exclusion of secondary infection sites, special handling of common commensals, MBI-LCBI logic, and surveillance adjudication that cannot be completely reconstructed from structured MIMIC-IV data [2]. The project therefore uses the term strict CVC-associated BSI proxy and reports broader and source-screened sensitivity labels.")

    add_heading(doc, "Why Precision-Recall Matters", 2)
    add_para(doc,
        "With event prevalence between roughly 1% and 5% depending on the prediction frame, a model can achieve a reasonable ROC-AUC while still producing many false positives. PR-AUC, PPV, recall, false reviews per true positive, and review capacity are consequently central. The project treats ROC-AUC as one component of evidence, not the headline by itself.")

    add_heading(doc, "Original Roadmap", 2)
    add_figure(doc, ROADMAP,
               "Figure 1. Original project roadmap. The work ultimately extended the planned static, time-series, modeling, and explainability phases with extensive leakage, outcome-validity, operational-policy, and external-feasibility audits.",
               width=6.0)
    page_break(doc)


def add_data_methods(doc):
    add_heading(doc, "2. Data, Cohort, and Outcome Design", 1)
    add_heading(doc, "Data Sources", 2)
    add_table(doc,
              ["Source", "Role in the project", "Important boundary"],
              [
                  ["MIMIC-IV", "Primary cohort, catheter procedures, labs, vitals, microbiology, therapies, diagnoses", "Single-center deidentified EHR; procedure absence does not prove line absence"],
                  ["ARMD-MGB", "External organism-rule and partial source-screen transportability", "Microbiology accession data; no catheter/predictor timeline"],
                  ["eICU-CRD", "External full-model feasibility assessment", "Insufficient microbiology-linked explicit catheter episodes"],
              ], [1.35, 2.65, 2.3], font_size=9.1, first_col_bold=True)

    add_heading(doc, "Catheter Episode Reconstruction", 2)
    add_para(doc,
        "The earliest versions selected a single eventual longest catheter per ICU stay. That choice uses future information and can attach cultures to the wrong exposure. Version 0.5 rebuilt the denominator from all recorded CVC procedure events, merged them into continuous exposure periods, retained every eligible period, and applied a minimum observed duration of 48 hours.")
    add_table(doc,
              ["Cohort quantity", "Count", "Interpretation"],
              [
                  ["Raw CVC procedure events", "26,285", "Positive documentation of catheter activity"],
                  ["Continuous exposure periods", "22,812", "Reconstructed patient-catheter exposure units"],
                  ["Eligible periods >=48 h", "11,602", "Primary denominator for later landmark modeling"],
                  ["Stays with multiple periods", "1,162", "Evidence that longest-line selection discarded meaningful structure"],
                  ["Broad proxy-positive episodes", "528", "Sensitive CVC-associated BSI proxy"],
                  ["Strict proxy-positive episodes", "371", "Organism-qualified stricter proxy"],
              ], [2.25, 1.0, 3.05], font_size=9.2, first_col_bold=True)

    add_heading(doc, "Daily Landmark Frame", 2)
    add_para(doc,
        "Run 17 converted eligible episodes into daily prediction rows. Each row represents what would have been knowable at a specific landmark, with a seven-day future outcome and explicit termination by line removal, procedure end, discharge, or death. This created 64,752 landmark rows before later source screening.")

    add_heading(doc, "Outcome Hierarchy", 2)
    add_table(doc,
              ["Label", "Logic", "Role"],
              [
                  ["Broad proxy", "Positive blood culture after >=48 h of observed line exposure", "Sensitivity analysis"],
                  ["Strict organism proxy", "Broad timing plus recognized pathogen or qualifying repeated common commensal", "Intermediate outcome"],
                  ["Primary-likely", "Strict proxy without structured evidence of another source", "Specific but sparse"],
                  ["Secondary-possible", "Strict proxy with plausible nearby nonblood source evidence", "Separated rather than silently counted as primary"],
                  ["Primary-or-uncertain", "Primary-likely plus cases lacking decisive source evidence", "Final seven-day modeling target"],
              ], [1.35, 3.3, 1.65], font_size=9.1, first_col_bold=True)

    add_heading(doc, "Feature Families", 2)
    add_para(doc,
        "The dynamic branch evaluated static context, rolling laboratory values, vital signs, trends, measurement recency, antibiotics, vasopressors, and care-process features. Measurement-intensity proxies were explicitly audited because test ordering can encode clinician concern rather than physiology. The final evidence package favors features that are available at the landmark and excludes the outcome-derived early-positive-culture flag.")

    add_heading(doc, "Temporal Evaluation", 2)
    add_table(doc,
              ["Period", "Role", "Handling"],
              [
                  ["2008-2013", "Model development", "Training data"],
                  ["2014-2016", "Calibration", "Platt calibration"],
                  ["2017-2019", "Leakage-safe characterization", "Primary final evidence reported in this report"],
                  ["2020-2022", "Historical temporal sensitivity", "Previously inspected; not treated as a pristine final lockbox"],
              ], [1.25, 2.05, 3.0], font_size=9.2, first_col_bold=True)
    add_callout(doc, "Evaluation principle",
                "Later, lower scores are more trustworthy because the reference time, feature availability, label definition, and cohort denominator were progressively made harder to exploit retrospectively.",
                fill=PALE_GOLD)
    page_break(doc)


def add_model_evolution(doc):
    add_heading(doc, "3. Model Evolution and Major Lessons", 1)
    add_para(doc,
        "The project should not be read as 33 attempts to tune the same classifier. It is a sequence of methodological corrections. Each phase changed what question was being answered and how closely the experiment resembled prospective use.")

    add_heading(doc, "Selected Milestones", 2)
    add_table(doc,
              ["Stage", "Representative result", "What was learned"],
              [
                  ["Runs 1-2: near-event static", "ROC-AUC 0.798-0.835; PR-AUC 0.237-0.292", "Static and rolling-lab data contained strong retrospective signal, but reference times depended on eventual outcomes."],
                  ["Runs 4.1-6: corrected strict static", "Run 5 ROC-AUC 0.789; PR-AUC 0.147", "Signal persisted under stricter organism logic, but calibration and alarm burden were poor."],
                  ["Runs 7-14: landmark dynamic", "Best v0.4H 168 h test ROC-AUC 0.706; PR-AUC 0.039", "Vitals and therapy context added some ranking value, but event rarity dominated and bedside alerting remained impractical."],
                  ["Runs 16-24: v0.5 redesign", "All episodes; daily landmarks; 7-day source-screened target", "Cohort definition and outcome validity mattered more than broad feature search."],
                  ["Runs 25-30: evaluation and leakage audit", "Final safe ROC-AUC 0.612; PR-AUC 0.065", "Removing outcome-derived information changed PR-AUC only modestly and preserved limited ranking signal."],
                  ["Runs 31-33: transportability and consolidation", "External label components supported; evidence package completed", "The remaining uncertainty is label validity and transportability, not another tuning sweep."],
              ], [1.55, 1.75, 3.0], font_size=9.0, first_col_bold=True)
    add_para(doc,
        "The numbers above are not directly comparable because labels, prevalence, episode construction, and evaluation periods changed. They illustrate the cost of methodological honesty, not a conventional performance leaderboard.",
        size=9, color=MUTED, italic=True, before=3, after=8)

    add_heading(doc, "Static Modeling: Useful but Retrospective", 2)
    add_para(doc,
        "The initial improved XGBoost model increased PR-AUC from 0.157 to 0.237, and static-plus-lab optimization reached PR-AUC 0.292. The apparent improvement was real within that retrospective task, but positive cases were characterized near culture while negatives were referenced near catheter removal. This made the model a near-event characterization tool rather than a deployable insertion-time baseline.")
    add_para(doc,
        "After correcting the audit feature and applying stricter organism logic, Run 5 achieved ROC-AUC 0.789 and PR-AUC 0.147. Run 6 showed why discrimination was not enough: at the automatically selected threshold, the model would flag 38.9 of every 100 stays, with PPV 5.4% and approximately 17.5 false alerts per true positive.")

    add_heading(doc, "Dynamic Modeling: Better Framing, Harder Task", 2)
    add_para(doc,
        "The landmark branch converted each episode into repeated prediction opportunities and tested 24-, 48-, 72-, and 168-hour horizons. Vitals, laboratory trajectories, therapy context, and care-process features were added in stages. The seven-day/168-hour horizon consistently offered the most plausible surveillance role because shorter horizons were extremely sparse and generated unstable precision.")
    add_para(doc,
        "Run 14 separated two functional roles: longer-horizon surveillance review and shorter-horizon targeted review. This was conceptually important even though the resulting PPV remained low. It established that one score should not be forced into incompatible operational workflows.")

    add_heading(doc, "v0.5: Cohort and Label Quality Before Complexity", 2)
    add_para(doc,
        "The v0.5 redesign retained all catheter exposure periods, standardized daily landmarks, used one seven-day outcome, incorporated competing events, constrained model families, and preserved temporal ordering. Subsequent runs then focused on source-screened labels, calibration, review budgets, locked error analysis, ICD agreement, and leakage audits rather than unrestricted model searching.")

    add_heading(doc, "The Leakage Finding", 2)
    add_para(doc,
        "Run 29 found that early_positive_culture encoded eventual culture positivity and could be present before the positive result was demonstrably available. The feature was removed from the final candidate. Validation PR-AUC fell by only 0.0038, suggesting that the broader model signal was not solely created by this defect.")
    add_figure(doc, FIGURES["leakage"],
               "Figure 2. Fraction of early-positive landmark rows for which an organism-positive microbiology result had not yet been stored. This audit motivated removal of the feature from the final safe candidate.",
               width=6.15)
    page_break(doc)


def add_final_evidence(doc):
    add_heading(doc, "4. Final Leakage-Safe Model Evidence", 1)
    add_heading(doc, "Frozen Candidate", 2)
    add_table(doc,
              ["Component", "Final specification"],
              [
                  ["Model", "XGBoost with Platt calibration"],
                  ["Prediction unit", "Daily landmark within an eligible CVC exposure episode"],
                  ["Horizon", "Strict primary-or-uncertain CVC-associated BSI proxy within 7 days"],
                  ["Evaluation population", "5,694 landmarks, 648 episodes, 590 patients"],
                  ["Positive burden", "242 positive landmarks (4.25%); 54 positive episodes"],
                  ["Leakage control", "early_positive_culture excluded"],
                  ["Intended analytical role", "Retrospective ranker for bounded infection-prevention review"],
              ], [2.0, 4.3], font_size=9.3, first_col_bold=True)

    add_heading(doc, "Discrimination, Calibration, and Uncertainty", 2)
    add_table(doc,
              ["Metric", "Estimate", "95% patient-clustered CI", "Interpretation"],
              [
                  ["ROC-AUC", "0.612", "0.518-0.703", "Modest discrimination; interval includes weak performance"],
                  ["PR-AUC", "0.065", "0.041-0.113", "Above 0.0425 prevalence; 1.54x lift"],
                  ["Brier score", "0.0405", "0.0299-0.0517", "Must be interpreted relative to prevalence"],
                  ["Brier Skill Score", "0.005", "-0.019 to 0.021", "No clear absolute-risk skill over prevalence baseline"],
                  ["Calibration slope", "0.807", "0.069-1.655", "Point estimate plausible, uncertainty very wide"],
                  ["Expected:observed", "1.159", "0.884-1.624", "Possible overprediction; interval remains broad"],
              ], [1.4, 0.85, 1.55, 2.5], font_size=9.0, first_col_bold=True)
    add_callout(doc, "Interpretation",
                "The model ranks some higher-risk episodes above lower-risk episodes, but its probability scale is not dependable enough for bedside absolute-risk communication.",
                fill=PALE_GOLD)

    add_figure(doc, FIGURES["calibration"],
               "Figure 3. Leakage-safe calibration by risk decile. The widening departure from the diagonal shows why the score should not be presented as a calibrated bedside probability.",
               width=6.0)

    add_heading(doc, "Bounded Review-List Performance", 2)
    add_table(doc,
              ["Review budget", "Episodes reviewed", "PPV", "Positive-episode recall", "False reviews / TP"],
              [
                  ["Top 1%", "7", "42.9%", "5.6%", "1.3"],
                  ["Top 2%", "13", "23.1%", "5.6%", "3.3"],
                  ["Top 5%", "33", "15.2%", "9.3%", "5.6"],
                  ["Top 10%", "65", "15.4%", "18.5%", "5.5"],
                  ["Top 20%", "130", "16.2%", "38.9%", "5.2"],
              ], [1.25, 1.2, 0.85, 1.45, 1.3], font_size=9.2, first_col_bold=True)
    add_para(doc,
        "The top-1% PPV is based on only seven reviewed episodes and has a 95% interval from 0% to 71.4%; it should not be treated as a stable operating point. The top-10% and top-20% budgets better show the likely workload-yield tradeoff.",
        size=9.3, color=MUTED, italic=True)
    add_figure(doc, FIGURES["policy"],
               "Figure 4. Episode-level review yield and workload across ranked review budgets. The score offers enrichment, but review capacity must be fixed in advance.",
               width=6.15)
    page_break(doc)


def add_clinical_interpretation(doc):
    add_heading(doc, "5. Clinical and Operational Interpretation", 1)
    add_heading(doc, "Where the Model Could Fit", 2)
    add_callout(doc, "Most defensible workflow",
                "A noninterruptive, bounded worklist reviewed by infection-prevention personnel or a research teamâ€”for example, the highest-ranked 5% to 10% of catheter episodes during a defined review period.",
                fill=PALE_GREEN)
    add_para(doc,
        "In this role, the score does not diagnose infection. It identifies records that may deserve earlier or more structured review. The final decision still depends on organism identity, line timing, alternative sources, symptoms, competing events, and clinical context.")

    add_heading(doc, "Where the Model Should Not Be Used", 2)
    add_table(doc,
              ["Unsafe use", "Why it is unsupported"],
              [
                  ["Interruptive nurse alert", "False-positive burden remains high and repeated landmarks can amplify alarm volume"],
                  ["Autonomous CLABSI diagnosis", "The endpoint is a proxy and lacks full NHSN source adjudication"],
                  ["Absolute-risk display", "Brier Skill Score is approximately zero and calibration uncertainty is wide"],
                  ["Line-removal recommendation", "No causal evidence that model-guided removal improves outcomes"],
                  ["Hospital comparison", "Single-center development and incomplete external model validation limit transportability"],
              ], [2.0, 4.3], font_size=9.2, first_col_bold=True, header_fill=DARK_BLUE)

    add_heading(doc, "How Bad Is the PPV?", 2)
    add_para(doc,
        "A PPV near 15% means roughly one true-positive episode for every six or seven reviewed episodes. That is poor for an interruptive clinical alarm but potentially workable for a deliberately capped retrospective review queue, especially if the alternative is unstructured manual surveillance. Suitability depends on reviewer capacity, the cost of missed cases, and whether the review action itself is low risk.")

    add_heading(doc, "Feature Interpretation", 2)
    add_para(doc,
        "Across the project, recurring signal came from line/site documentation, platelet behavior, white blood cell trajectories, age, lactate, creatinine, catheter type, vital signs, and therapy context. Several of these features may reflect measurement intensity or clinician concern as much as underlying pathophysiology. The project therefore treats feature importance as descriptive model behavior, not causal evidence.")

    add_heading(doc, "Subgroup Findings", 2)
    add_para(doc,
        "Run 30 evaluated prespecified subgroup cells with patient-clustered intervals. Only seven levels met the more stable reporting rule of at least 100 patients and 20 positive patients; eleven additional levels required cautious interpretation. Intervals overlapped widely, so the analysis does not establish fairness, harm, or causal differences between groups.")
    add_figure(doc, FIGURES["subgroups"],
               "Figure 5. PR-AUC lift by prespecified subgroup. Blue cells met the more-stable reporting rule; orange cells require cautious interpretation.",
               width=6.1)
    page_break(doc)


def add_external_evidence(doc):
    add_heading(doc, "6. External Evidence", 1)
    add_heading(doc, "External Validation Feasibility", 2)
    add_para(doc,
        "The project tested two independent databases before forcing an external performance estimate. ARMD-MGB contained rich microbiology but no catheter exposure or longitudinal predictor timeline. eICU contained multicenter ICU data, but only one positive-blood-culture stay met the explicit >=48-hour placement criterion under the attempted mapping. Neither database could honestly reproduce the frozen seven-day catheter-episode task.")
    add_table(doc,
              ["Database", "What could be tested", "What could not be tested", "Conclusion"],
              [
                  ["ARMD-MGB", "Organism rule; urine/respiratory source-window sensitivity", "Catheter episodes, landmarks, model performance", "External label-component evidence"],
                  ["eICU-CRD", "Availability and linkage feasibility", "Reliable full cohort and outcome reconstruction", "Failed feasibility gate"],
              ], [1.15, 2.0, 2.1, 1.05], font_size=9.0, first_col_bold=True)

    add_heading(doc, "ARMD-MGB Label-Component Findings", 2)
    add_para(doc,
        "Run 32 characterized 32,887 positive blood-culture accessions. After harmonizing organism names, 28 organism categories were shared with MIMIC-IV. Rank correlation was 0.651, while Jensen-Shannon divergence was 0.456, indicating moderate ordering agreement alongside meaningful case-mix shift. This supports portability of the organism logic, not portability of the prediction model.")
    add_figure(doc, FIGURES["organisms"],
               "Figure 6. Relative organism profiles among qualifying MIMIC-IV episodes and ARMD-MGB strict-rule accessions after canonical name harmonization.",
               width=6.15)
    add_figure(doc, FIGURES["source_window"],
               "Figure 7. Same-organism urine or respiratory evidence across prespecified ARMD-MGB source-search windows. These matches suggest possible secondary sources but do not establish NHSN attribution.",
               width=6.1)

    add_callout(doc, "Evidence boundary",
                "External testing supports only the portability of selected organism and source-screening components. No independent database reproduced the full catheter-episode prediction task, so no external model-performance claim is made.",
                fill=PALE_GOLD)
    page_break(doc)


def add_current_state(doc):
    add_heading(doc, "7. Current Project State", 1)
    add_heading(doc, "What Is in Good Shape", 2)
    for text in [
        "A reproducible end-to-end pipeline exists from raw MIMIC tables through catheter episode construction, landmark features, modeling, calibration, and evaluation.",
        "The largest cohort-design vulnerabilityâ€”selecting the eventual longest lineâ€”was removed in v0.5.",
        "The final candidate excludes a confirmed outcome-derived feature and has patient-clustered uncertainty estimates.",
        "The outcome terminology is appropriately conservative and separates likely primary, possible secondary, and uncertain source patterns.",
        "Operational interpretation includes review budgets, PPV, recall, and false reviews per true positive.",
        "External database feasibility was evaluated honestly, including a documented failed feasibility gate.",
        "Run 33 provides a consolidated, traceable endpoint with an explicit claim register and evidence boundaries.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "What Still Needs Improvement", 2)
    add_table(doc,
              ["Domain", "Remaining issue", "Most useful next work"],
              [
                  ["Outcome validity", "Proxy rather than adjudicated CLABSI", "If resumed, conduct independent blinded clinical review"],
                  ["Line denominator", "procedureevents absence does not prove no line", "Validate against a richer catheter registry or EHR"],
                  ["External performance", "No independent cohort reproduces the full task", "Use longitudinal EHR with line timing, microbiology, and daily predictors"],
                  ["Calibration", "Near-zero Brier skill and wide intervals", "External recalibration only after transportable discrimination is demonstrated"],
                  ["Review burden", "Most flagged episodes are false positives", "Prespecify bounded review capacity and evaluate workflow value"],
                  ["Subgroups", "Sparse positives and wide uncertainty", "Larger external sample and prespecified subgroup analysis"],
              ], [1.3, 2.45, 2.55], font_size=9.1, first_col_bold=True)

    add_heading(doc, "What the Project Demonstrates", 2)
    add_callout(doc, "Technical achievement",
                "The project demonstrates senior-level clinical ML practice: cohort redesign, temporal leakage detection, rare-event evaluation, dynamic landmark construction, calibration analysis, operational burden translation, label sensitivity analysis, external feasibility gates, and honest limitation management.",
                fill=PALE_BLUE)
    add_para(doc,
        "The strongest contribution is not a single performance number. It is the conversion of an initially optimistic static model into a more defensible research system whose intended use and uncertainty are explicit.")

    page_break(doc)
    add_heading(doc, "Future Work if the Project Is Resumed", 2)
    add_table(doc,
              ["Priority", "Action", "Decision enabled"],
              [
                  ["1", "Validate the proxy through independent blinded clinical review", "Estimate label agreement and identify failure modes"],
                  ["2", "Continue searching for a longitudinal external EHR", "Determine whether discrimination and review yield transport across institutions"],
                  ["3", "Keep broad tuning frozen on repeatedly inspected MIMIC periods", "Protect the interpretability of future evidence"],
                  ["4", "If new data arrive, prespecify mapping and review budgets before outcomes", "Reduce researcher degrees of freedom"],
              ], [0.65, 3.55, 2.1], font_size=9.1, first_col_bold=True)
    add_heading(doc, "A 60-Second Explanation", 2)
    add_callout(doc, "Presentation version",
                "CVCML began as a high-performing static XGBoost model for catheter-associated bloodstream infection risk. The strongest early scores depended on retrospective timing, so we rebuilt the cohort around all line episodes, daily landmarks, and a seven-day source-screened outcome. The final leakage-safe model has modest discrimination and weak calibration, but it can enrich a bounded infection-prevention review list. Run 33 is the completed project endpoint; any future extension should focus on independent label validation or external transportability, not more tuning.",
                fill=PALE_BLUE)
    add_heading(doc, "Three Takeaways to Emphasize", 2)
    for text in [
        "The project improved scientifically even when the headline metrics decreased.",
        "The final model is a review-prioritization research tool, not an alarm or diagnosis.",
        "The remaining uncertainty is mainly outcome validity and transportability, not a missing hyperparameter search.",
    ]:
        add_bullet(doc, text)
    page_break(doc)


RUNS = [
    ("1", "Static + improved baseline", "Built baseline and static-plus-lab XGBoost models; established that rolling laboratory features improved retrospective discrimination."),
    ("2", "Static optimization", "Compared logistic regression, random forest, static XGBoost, tuned XGBoost, and calibration; tuning did not beat the untuned static-plus-lab PR-AUC."),
    ("3", "Leakage robustness", "Ablated total dwell and site documentation; exposed strong outcome/reference-time dependence."),
    ("4", "Corrected static v0.2", "Rebuilt timing and strict culture rules but accidentally included an audit-derived feature."),
    ("4.1", "No-audit-leakage correction", "Removed early_positive_culture; became the cleanest pragmatic-label static benchmark at that stage."),
    ("5", "Strict organism sensitivity", "Required pathogen or repeated-commensal logic; established the strongest sensitivity-tested static characterization."),
    ("6", "Static characterization", "Added bootstrap intervals, calibration, subgroup checks, and explicit alert burden."),
    ("7", "First landmark dynamic model", "Created repeated landmarks and future-event horizons; showed severe rare-event limitations."),
    ("8", "Vitals enrichment", "Added charted physiology and horizon sensitivity; 168-hour surveillance was most promising."),
    ("9", "Proxy robustness", "Removed measurement-intensity proxies and tested gray-zone labels; physiologic value signal persisted."),
    ("10", "Therapy context", "Added antibiotics and vasopressors; therapy context produced modest gains."),
    ("11", "Cleaned therapy context", "Separated clinically meaningful therapy features from medication-order artifacts."),
    ("12", "Dynamic alert policy", "Tested calibration, cooldowns, caps, and review policies; calibration did not solve low PPV."),
    ("13", "Care-process context", "Added caregiver and line-care-process signals; limited incremental value."),
    ("14", "Split dynamic use cases", "Separated longer-horizon surveillance from shorter-horizon targeted review."),
    ("15", "Candidate characterization", "Compared static and dynamic candidates by functional role and preserved parallel interpretations."),
    ("16", "v0.5 episode redesign", "Reconstructed all catheter exposure periods and removed eventual-longest-line selection."),
    ("17", "Daily landmark frame", "Created one seven-day prediction target with explicit competing/censoring outcomes."),
    ("18", "Constrained development models", "Compared logistic regression and XGBoost feature sets while keeping the later period sealed."),
    ("19", "Calibration and review policy", "Evaluated Platt calibration and capacity-limited top-risk review."),
    ("20", "Dynamic feature enrichment", "Added raw vitals, antibiotics, and vasopressors to the redesigned v0.5 frame."),
    ("21", "Target review framing", "Compared short-term horizons and episode-level review yield."),
    ("22", "Secondary-source label audit", "Separated primary-likely, secondary-possible, and uncertain source patterns."),
    ("23", "Label sensitivity modeling", "Held models fixed while comparing original strict, primary-likely, and primary-or-uncertain targets."),
    ("24", "Operating policy", "Characterized row and episode review budgets for the selected source-screened target."),
    ("25", "Temporal evaluation", "Opened the 2020-2022 period once for the then-frozen candidate; later treated as historical sensitivity after further audits."),
    ("26", "Locked error analysis", "Characterized false positives, missed positives, episode capture, and organism patterns without retraining."),
    ("27", "ICD agreement", "Measured weak agreement between structured proxy labels and ICD catheter-infection codes."),
    ("27.1", "Discordance supplement", "Stratified proxy/ICD failure modes and created a balanced manual-review sample."),
    ("28", "Model card and results package", "Consolidated the then-current model definition, performance, and limitations."),
    ("29", "Outcome validity and leakage audit", "Detected the early_positive_culture prospectivity defect, removed it, and created an adjudication protocol."),
    ("30", "Safe candidate characterization", "Produced patient-clustered metric intervals, Brier Skill Score, subgroup checks, calibration, and review-burden estimates."),
    ("31", "External feasibility", "Audited ARMD-MGB and eICU before external modeling; neither could reproduce the full frozen task."),
    ("32", "External label transportability", "Validated organism and partial source-screen components in ARMD-MGB and quantified distribution shift."),
    ("33", "Evidence consolidation", "Consolidated leakage-safe results, external evidence boundaries, and a claim register."),
]


def add_appendix(doc):
    add_heading(doc, "Appendix A. Run-by-Run Record", 1)
    add_para(doc,
        "This compact index preserves the complete project arc. Runs often changed the cohort, outcome, or evaluation question, so they should not be interpreted as repeated tuning on one fixed benchmark.",
        size=9.5, color=MUTED)
    add_table(doc, ["Run", "Focus", "Outcome / lesson"], RUNS,
              [0.55, 1.55, 4.2], font_size=8.5, first_col_bold=True, zebra=True)

    add_heading(doc, "Appendix B. Metric Glossary", 1)
    add_table(doc,
              ["Metric", "Meaning in this project"],
              [
                  ["ROC-AUC", "Probability that a randomly selected positive row ranks above a randomly selected negative row."],
                  ["PR-AUC", "Average precision across thresholds; more informative than ROC-AUC when positives are rare."],
                  ["PR-AUC lift", "PR-AUC divided by event prevalence; values above 1 indicate enrichment over random ranking."],
                  ["PPV / precision", "Proportion of flagged episodes that are positive under the project proxy."],
                  ["Recall / sensitivity", "Proportion of positive episodes captured by a threshold or review budget."],
                  ["Brier score", "Mean squared error of predicted probabilities; must be compared with a prevalence baseline."],
                  ["Brier Skill Score", "Relative improvement over the constant-prevalence predictor; 0 means no improvement."],
                  ["Calibration slope", "Whether predictions vary too much or too little; ideal value is 1."],
                  ["Expected:observed ratio", "Total predicted events divided by observed events; ideal value is 1."],
                  ["False reviews / TP", "Number of negative episodes reviewed for each true-positive episode captured."],
              ], [1.6, 4.7], font_size=9.0, first_col_bold=True)

    add_heading(doc, "Appendix C. Selected References", 1)
    references = [
        "[1] Johnson AEW, Bulgarelli L, Shen L, et al. MIMIC-IV, a freely accessible electronic health record dataset. Scientific Data. 2023;10:1. doi:10.1038/s41597-022-01899-x.",
        "[2] Centers for Disease Control and Prevention. NHSN Patient Safety Component Manual, Chapter 4: Bloodstream Infection Event (Central Line-Associated Bloodstream Infection and Non-central Line-Associated Bloodstream Infection). January 2026.",
        "[3] Rahmani K, et al. Early prediction of central line associated bloodstream infection using machine learning. American Journal of Infection Control. 2022;50. XGBoost AUROC reported as 0.762.",
        "[4] Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378.",
        "[5] Wei Z, Kanjilal S. Antibiotic Resistance Microbiology Dataset Mass General Brigham (ARMD-MGB), version 1.0.0. PhysioNet. 2025. doi:10.13026/2r5k-b955.",
        "[6] Pollard TJ, Johnson AEW, Raffa JD, Celi LA, Badawi O, Mark RG. eICU Collaborative Research Database, version 2.0. PhysioNet. doi:10.13026/C2WM1R.",
    ]
    for ref in references:
        p = add_para(doc, ref, size=8.5, color=INK, after=3, line=1.0)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    add_heading(doc, "Appendix D. Project File Map", 1)
    for location, purpose in [
        ("SRC/", "Versioned extraction, feature engineering, modeling, and audit scripts"),
        ("Outputs/Run 1 ... Run 33/", "Run-specific metrics, plots, models, manifests, and notes"),
        ("Notebooks/CVCML Project Hub - updated v0.5 Run33.docx", "Detailed chronological project log"),
        ("Outputs/Run 33 (v0.5 Publication Consolidation)/", "Current consolidated model evidence and figures"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        set_run(p.add_run(f"{location}: "), size=8.8, bold=True, color=DARK_BLUE)
        set_run(p.add_run(purpose), size=8.8, color=INK)

    add_para(doc, "End of report", size=9, bold=True, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=0)


def build():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_exec_summary(doc)
    add_clinical_premise(doc)
    add_data_methods(doc)
    add_model_evolution(doc)
    add_final_evidence(doc)
    add_clinical_interpretation(doc)
    add_external_evidence(doc)
    add_current_state(doc)
    add_appendix(doc)
    doc.core_properties.title = "CVCML Comprehensive Project Report - Runs 1-33"
    doc.core_properties.subject = "Leakage-audited CVC-associated bloodstream infection risk modeling in MIMIC-IV"
    doc.core_properties.author = "CVCML Project"
    doc.core_properties.keywords = "MIMIC-IV, CVC, bloodstream infection, XGBoost, clinical machine learning, landmark modeling"
    doc.save(REPORT_DOCX)
    print(REPORT_DOCX)


if __name__ == "__main__":
    build()

